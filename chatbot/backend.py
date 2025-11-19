from bs4 import BeautifulSoup
from datetime import datetime, timedelta, timezone
import time, json, os, re, random, feedparser, requests, tempfile, pdfkit, base64, uuid, shutil, cv2, yt_dlp, urllib.parse
from docx import Document
from docx.shared import Inches
from typing import List, Dict, Any
from pytube import YouTube
from youtube_transcript_api import YouTubeTranscriptApi
import numpy as np
from moviepy import VideoFileClip
from urllib.parse import urlparse

# ==============================================
# LỚP QUẢN LÝ DỮ LIỆU - ĐÃ SỬA ĐỂ KẾT NỐI VỚI ADMIN
# ==============================================

class DataManager:
    def __init__(self):
        self.rss_feeds_file = "rss_feeds.json"
        self.youtube_channels_file = "youtube_channels.json"
        self._ensure_data_files()

    def _ensure_data_files(self):
        """Đảm bảo file dữ liệu tồn tại, nếu không thì tạo từ admin"""
        try:
            # Nếu file không tồn tại, tạo từ dữ liệu mặc định của admin
            if not os.path.exists(self.rss_feeds_file):
                print("🔄 Không tìm thấy file RSS feeds, đang tạo từ admin...")
                self._create_default_rss_feeds()

            if not os.path.exists(self.youtube_channels_file):
                print("🔄 Không tìm thấy file YouTube channels, đang tạo từ admin...")
                self._create_default_youtube_channels()

        except Exception as e:
            print(f"❌ Lỗi khởi tạo file dữ liệu: {e}")

    def _create_default_rss_feeds(self):
        """Tạo RSS feeds mặc định từ admin"""
        default_feeds = {
            'thời sự': 'https://video.vnexpress.net/rss/thoi-su.rss',
            'kinh doanh': 'https://video.vnexpress.net/rss/kinh-doanh.rss',
            'công nghệ': 'https://video.vnexpress.net/rss/khoa-hoc-cong-nghe.rss',
            'thể thao': 'https://video.vnexpress.net/rss/the-thao.rss',
            'giáo dục': 'https://video.vnexpress.net/rss/giao-duc.rss',
            'sức khỏe': 'https://video.vnexpress.net/rss/suc-khoe.rss',
            'đời sống': 'https://video.vnexpress.net/rss/nhip-song.rss',
            'giải trí': 'https://video.vnexpress.net/rss/giai-tri.rss',
            'du lịch': 'https://video.vnexpress.net/rss/du-lich.rss',
            'pháp luật': 'https://video.vnexpress.net/rss/phap-luat.rss',
            'video thời sự': 'https://thanhnien.vn/rss/video/thoi-su.rss'
        }
        self.save_rss_feeds(default_feeds)

    def _create_default_youtube_channels(self):
        """Tạo YouTube channels mặc định từ admin"""
        default_channels = {
            'VTV24': 'UCabsTV34JwALXKGMqHpvUiA',
            'TIN TỨC VIỆT': 'UCxlprA9Y_T98gDqXMc46baw',
            'Tin24h': 'UCUmRGR3a-g13O6pG927KQmg'
        }
        self.save_youtube_channels(default_channels)

    def load_rss_feeds(self):
        """Tải RSS feeds từ file - KẾT NỐI TRỰC TIẾP VỚI ADMIN"""
        try:
            if os.path.exists(self.rss_feeds_file):
                with open(self.rss_feeds_file, 'r', encoding='utf-8') as f:
                    feeds = json.load(f)
                    print(f"✅ Đã tải {len(feeds)} RSS feeds từ admin")
                    return feeds
            else:
                print("❌ Không tìm thấy file RSS feeds từ admin")
                return self._create_default_rss_feeds()
        except Exception as e:
            print(f"❌ Lỗi tải RSS feeds từ admin: {e}")
            return self._create_default_rss_feeds()

    def load_youtube_channels(self):
        """Tải YouTube channels từ file - KẾT NỐI TRỰC TIẾP VỚI ADMIN"""
        try:
            if os.path.exists(self.youtube_channels_file):
                with open(self.youtube_channels_file, 'r', encoding='utf-8') as f:
                    channels = json.load(f)
                    print(f"✅ Đã tải {len(channels)} YouTube channels từ admin")
                    return channels
            else:
                print("❌ Không tìm thấy file YouTube channels từ admin")
                return self._create_default_youtube_channels()
        except Exception as e:
            print(f"❌ Lỗi tải YouTube channels từ admin: {e}")
            return self._create_default_youtube_channels()

    def save_rss_feeds(self, feeds):
        """Lưu RSS feeds vào file - DÙNG CHUNG VỚI ADMIN"""
        try:
            with open(self.rss_feeds_file, 'w', encoding='utf-8') as f:
                json.dump(feeds, f, ensure_ascii=False, indent=2)
            print(f"💾 Đã lưu {len(feeds)} RSS feeds (đồng bộ với admin)")
            return True
        except Exception as e:
            print(f"❌ Lỗi lưu RSS feeds: {e}")
            return False

    def save_youtube_channels(self, channels):
        """Lưu YouTube channels vào file - DÙNG CHUNG VỚI ADMIN"""
        try:
            with open(self.youtube_channels_file, 'w', encoding='utf-8') as f:
                json.dump(channels, f, ensure_ascii=False, indent=2)
            print(f"💾 Đã lưu {len(channels)} YouTube channels (đồng bộ với admin)")
            return True
        except Exception as e:
            print(f"❌ Lỗi lưu YouTube channels: {e}")
            return False

    def refresh_data(self):
        """Làm mới dữ liệu từ admin - GỌI KHI CẦN CẬP NHẬT"""
        print("🔄 Đang làm mới dữ liệu từ admin...")
        rss_feeds = self.load_rss_feeds()
        youtube_channels = self.load_youtube_channels()
        return rss_feeds, youtube_channels

# ==============================================
# LỚP QUẢN LÝ BACKUP - MỚI THÊM
# ==============================================

class BackupManager:
    def __init__(self, history_file="chat_history.json"):
        self.history_file = history_file
        self.backup_dir = "chat_backups"
        os.makedirs(self.backup_dir, exist_ok=True)

    def create_backup(self):
        """Tạo bản backup tự động"""
        if not os.path.exists(self.history_file):
            return

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_file = os.path.join(self.backup_dir, f"chat_backup_{timestamp}.json")

        try:
            shutil.copy2(self.history_file, backup_file)
            print(f"✅ Đã tạo backup: {backup_file}")

            # Giữ chỉ 10 backup gần nhất
            self._clean_old_backups()
        except Exception as e:
            print(f"❌ Lỗi tạo backup: {e}")

    def _clean_old_backups(self):
        """Chỉ giữ 10 backup gần nhất"""
        try:
            backups = []
            for file in os.listdir(self.backup_dir):
                if file.startswith("chat_backup_") and file.endswith(".json"):
                    file_path = os.path.join(self.backup_dir, file)
                    backups.append((file_path, os.path.getctime(file_path)))

            # Sắp xếp theo thời gian tạo
            backups.sort(key=lambda x: x[1], reverse=True)

            # Xóa các backup cũ
            for backup in backups[10:]:
                try:
                    os.remove(backup[0])
                    print(f"🗑️ Đã xóa backup cũ: {backup[0]}")
                except Exception as e:
                    print(f"❌ Lỗi xóa backup: {e}")
        except Exception as e:
            print(f"❌ Lỗi dọn dẹp backup: {e}")

# ==============================================
# LỚP QUẢN LÝ LỊCH SỬ CHAT - MỚI THÊM
# ==============================================

class ChatHistoryManager:
    def __init__(self, storage_file="chat_history.json"):
        self.storage_file = storage_file
        self.chat_sessions = []
        self.current_session_id = None
        self.max_sessions = 50
        self.max_days = 40
        self.backup_manager = BackupManager(storage_file)

        # Đảm bảo thư mục tồn tại
        os.makedirs(os.path.dirname(os.path.abspath(self.storage_file)), exist_ok=True)

        # Tải lịch sử khi khởi động
        self.load_history()

        # Tạo session mới nếu chưa có
        if not self.current_session_id or not self.get_session(self.current_session_id):
            self.create_new_session("Phiên làm việc mới")
        else:
            print(f"🔄 Đã khôi phục phiên hiện tại: {self.current_session_id}")

    def get_vietnam_time(self):
        """Lấy thời gian hiện tại theo múi giờ Việt Nam (UTC+7)"""
        utc_now = datetime.now(timezone.utc)
        vietnam_time = utc_now + timedelta(hours=7)
        return vietnam_time

    def format_vietnam_time(self, dt=None):
        """Định dạng thời gian Việt Nam"""
        if dt is None:
            dt = self.get_vietnam_time()
        return dt.strftime('%d/%m/%Y %H:%M:%S')

    def ensure_data_integrity(self):
        """Đảm bảo tính toàn vẹn dữ liệu"""
        try:
            valid_sessions = []
            for session in self.chat_sessions:
                try:
                    # Đảm bảo các trường bắt buộc tồn tại
                    if 'id' not in session:
                        session['id'] = str(uuid.uuid4())
                    if 'title' not in session:
                        session['title'] = "Phiên không có tiêu đề"
                    if 'messages' not in session:
                        session['messages'] = []
                    if 'message_count' not in session:
                        session['message_count'] = len(session['messages'])
                    if 'created_at' not in session:
                        session['created_at'] = self.format_vietnam_time()
                    if 'updated_at' not in session:
                        session['updated_at'] = self.format_vietnam_time()

                    # Kiểm tra định dạng messages
                    valid_messages = []
                    for msg in session['messages']:
                        if isinstance(msg, dict) and 'role' in msg and 'content' in msg:
                            valid_messages.append(msg)

                    session['messages'] = valid_messages
                    session['message_count'] = len(valid_messages)
                    valid_sessions.append(session)

                except Exception as e:
                    print(f"⚠️ Sửa lỗi session: {e}")
                    continue

            self.chat_sessions = valid_sessions

            # Lưu lại sau khi sửa
            self.save_history()
            print("✅ Đã kiểm tra tính toàn vẹn dữ liệu lịch sử")
        except Exception as e:
            print(f"❌ Lỗi kiểm tra tính toàn vẹn: {e}")

    def load_history(self):
        """Tải lịch sử từ file"""
        try:
            if os.path.exists(self.storage_file):
                with open(self.storage_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.chat_sessions = data.get('chat_sessions', [])
                    self.current_session_id = data.get('current_session_id')

                print(f"✅ Đã tải {len(self.chat_sessions)} phiên chat từ lịch sử")

                # Kiểm tra tính toàn vẹn dữ liệu
                self.ensure_data_integrity()

                # Lọc các phiên cũ hơn 40 ngày
                self._clean_old_sessions()

            else:
                print("ℹ️ Chưa có file lịch sử, bắt đầu mới")
                self.chat_sessions = []
        except Exception as e:
            print(f"❌ Lỗi khi tải lịch sử: {e}")
            # Khởi tạo mới nếu có lỗi
            self.chat_sessions = []
            self.current_session_id = None

    def save_history(self):
        """Lưu lịch sử vào file - CÓ BACKUP"""
        try:
            # Tạo backup trước khi lưu
            self.backup_manager.create_backup()

            data = {
                'chat_sessions': self.chat_sessions,
                'current_session_id': self.current_session_id,
                'last_updated': self.format_vietnam_time(),
                'version': '1.0'
            }

            # Lưu tạm thời vào file tạm trước
            temp_file = self.storage_file + ".tmp"
            with open(temp_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            # Thay thế file cũ
            if os.path.exists(self.storage_file):
                os.remove(self.storage_file)
            os.rename(temp_file, self.storage_file)

            print(f"💾 Đã lưu lịch sử: {len(self.chat_sessions)} phiên")
        except Exception as e:
            print(f"❌ Lỗi khi lưu lịch sử: {e}")

    def _clean_old_sessions(self):
        """Xóa các phiên chat cũ hơn 40 ngày"""
        if not self.chat_sessions:
            return

        cutoff_date = self.get_vietnam_time() - timedelta(days=self.max_days)
        initial_count = len(self.chat_sessions)

        valid_sessions = []
        for session in self.chat_sessions:
            try:
                session_date = self._parse_datetime(session['updated_at'])
                if session_date > cutoff_date:
                    valid_sessions.append(session)
            except:
                # Giữ lại session nếu không parse được date
                valid_sessions.append(session)

        removed_count = initial_count - len(valid_sessions)
        self.chat_sessions = valid_sessions

        if removed_count > 0:
            print(f"🗑️ Đã xóa {removed_count} phiên chat cũ hơn {self.max_days} ngày")
            self.save_history()

    def _parse_datetime(self, date_string):
        """Chuyển chuỗi ngày tháng thành đối tượng datetime"""
        try:
            return datetime.strptime(date_string, '%d/%m/%Y %H:%M:%S')
        except:
            return self.get_vietnam_time()

    def create_new_session(self, title="Cuộc trò chuyện mới"):
        """Tạo session chat mới"""
        session_id = str(uuid.uuid4())
        current_time = self.format_vietnam_time()

        session = {
            'id': session_id,
            'title': title,
            'created_at': current_time,
            'updated_at': current_time,
            'messages': [],
            'message_count': 0
        }

        # Giới hạn số lượng session
        if len(self.chat_sessions) >= self.max_sessions:
            self.chat_sessions.pop(0)

        self.chat_sessions.append(session)
        self.current_session_id = session_id

        # Lưu ngay sau khi tạo
        self.save_history()
        print(f"🆕 Đã tạo phiên mới: {title}")
        return session_id

    def add_message(self, session_id: str, role: str, content: str, metadata: Dict = None):
        """Thêm tin nhắn vào session"""
        session = self.get_session(session_id)
        if session:
            current_time = self.format_vietnam_time()

            message = {
                'id': str(uuid.uuid4()),
                'role': role,
                'content': content,
                'timestamp': current_time,
                'metadata': metadata or {}
            }
            session['messages'].append(message)
            session['message_count'] = len(session['messages'])
            session['updated_at'] = current_time

            # Cập nhật tiêu đề nếu là tin nhắn đầu tiên
            if len(session['messages']) == 1 and role == 'user':
                session['title'] = content[:30] + "..." if len(content) > 30 else content

            # Lưu ngay sau khi thêm tin nhắn
            self.save_history()

    def get_session(self, session_id: str) -> Dict:
        """Lấy session theo ID"""
        for session in self.chat_sessions:
            if session['id'] == session_id:
                return session
        return None

    def get_all_sessions(self) -> List[Dict]:
        """Lấy tất cả session (sắp xếp mới nhất trước)"""
        return sorted(self.chat_sessions,
                     key=lambda x: self._parse_datetime(x['updated_at']),
                     reverse=True)

    def delete_session(self, session_id: str):
        """Xóa session"""
        self.chat_sessions = [s for s in self.chat_sessions if s['id'] != session_id]
        if self.current_session_id == session_id:
            if self.chat_sessions:
                self.current_session_id = self.chat_sessions[0]['id']
            else:
                self.current_session_id = None
        self.save_history()

    def clear_all_sessions(self):
        """Xóa toàn bộ lịch sử"""
        self.chat_sessions = []
        self.current_session_id = None
        self.save_history()

    def get_current_session_messages(self):
        """Lấy tin nhắn của session hiện tại"""
        if not self.current_session_id:
            return []
        session = self.get_session(self.current_session_id)
        return session['messages'] if session else []

# ==============================================
# CẤU HÌNH NÂNG CAO - ĐÃ CẬP NHẬT ĐỂ DÙNG NGUỒN TỪ ADMIN
# ==============================================

# KHÔNG DÙNG DEFAULT CỨNG NỮA - SẼ LOAD TỪ ADMIN

CATEGORY_ICONS = {
    'thời sự': '📰', 'kinh doanh': '💼', 'thể thao': '⚽', 'công nghệ': '🤖',
    'tin tức youtube': '📹', 'giải trí': '🎬', 'pháp luật': '⚖️', 'thế giới': '🌍',
    'công nghệ youtube': '💻', 'thể thao youtube': '🏀', 'video thời sự': '🎥',
    'giáo dục': '🎓', 'giáo dục youtube': '📚', 'thanh niên tổng hợp': '📺',
    'bí quyết ôn thi': '📚', 'phim ngắn': '🎭', 'thời sự thanh niên': '📡',
    'đời sống thanh niên': '🏠', 'giải trí thanh niên': '🎪', 'nld tổng hợp': '📰',
    'bản tin nld': '📋', 'thời sự trong nước nld': '🇻🇳', 'thời sự quốc tế nld': '🌏',
    'diễn đàn kinh tế': '💹', 'người đưa tin tổng hợp': '📢', 'hình sự': '🚔',
    'mới nóng': '🔥', 'giải trí người đưa tin': '🎭', 'xa lộ': '🚗', 'video xe': '🚙',
    'tin mới tổng hợp': '🆕', 'thời sự vnexpress': '📡', 'kinh doanh vnexpress': '💼',
    'công nghệ vnexpress': '💻', 'thể thao vnexpress': '⚽', 'giáo dục vnexpress': '🎓',
    'sức khỏe vnexpress': '🏥', 'đời sống vnexpress': '🏠', 'giải trí vnexpress': '🎬',
    'du lịch vnexpress': '✈️', 'pháp luật vnexpress': '⚖️', 'VTV24': '📺',
    'PDT youtube - Tin tức trong ngày': '📰', 'TIN TỬC VIỆT': '🇻🇳',
    'Tin24h': '🕐', 'Tin360 News': '🔄', 'công nghệ': '🤖', 'thực tế ảo': '🥽',
    'ai': '🧠', 'trí tuệ nhân tạo': '🤖'
}

# ==============================================
# HỆ THỐNG TỪ KHÓA MỞ RỘNG - MỚI THÊM
# ==============================================

# Hệ thống từ khóa mở rộng với ít nhất 45 từ khóa cho mỗi kênh/chủ đề
EXPANDED_KEYWORDS = {
    # Công nghệ - AI - Robot (45+ từ khóa) - ĐÃ BỔ SUNG TỪ KHÓA MỚI
    'công nghệ': [
        'ai', 'trí tuệ nhân tạo', 'artificial intelligence', 'machine learning',
        'robot', 'robotics', 'automation', 'tự động hóa', 'công nghệ số',
        'chuyển đổi số', 'digital transformation', 'blockchain', 'crypto',
        'bitcoin', 'metaverse', 'vũ trụ ảo', 'iot', 'internet of things',
        'big data', 'dữ liệu lớn', 'cloud computing', 'điện toán đám mây',
        'cybersecurity', 'an ninh mạng', 'hacking', 'bảo mật', 'privacy',
        'quantum computing', 'máy tính lượng tử', 'vr', 'ar', 'virtual reality',
        'augmented reality', 'thực tế ảo', 'thực tế tăng cường', '5g', '6g',
        'smartphone', 'điện thoại thông minh', 'tablet', 'máy tính bảng',
        'laptop', 'máy tính xách tay', 'pc', 'máy tính để bàn', 'gaming',
        'trò chơi điện tử', 'esports', 'thiết bị thông minh', 'smart device',
        'app', 'ứng dụng', 'software', 'phần mềm', 'hardware', 'phần cứng',
        'chip', 'vi xử lý', 'processor', 'samsung', 'apple', 'google',
        'microsoft', 'facebook', 'amazon', 'tesla', 'spacex', 'elon musk',
        # THÊM TỪ KHÓA MỚI ĐỂ ĐẢM BẢO TÌM KIẾM TẤT CẢ CHỦ ĐỀ
        'thực tế ảo', 'vr', 'virtual reality', 'công nghệ thực tế ảo',
        'công nghệ mới', 'đổi mới công nghệ', 'sáng tạo công nghệ'
    ],

    'công nghệ vnexpress': [
        'ai', 'trí tuệ nhân tạo', 'công nghệ mới', 'đổi mới sáng tạo',
        'startup công nghệ', 'công nghệ việt nam', 'phát minh', 'sáng chế',
        'research and development', 'r&d', 'lab', 'phòng thí nghiệm',
        'tech news', 'tin công nghệ', 'review công nghệ', 'đánh giá sản phẩm',
        'smartphone mới', 'laptop mới', 'tablet mới', 'wearable', 'đeo được',
        'smartwatch', 'đồng hồ thông minh', 'fitness tracker', 'theo dõi sức khỏe',
        'công nghệ y tế', 'health tech', 'edtech', 'công nghệ giáo dục',
        'fintech', 'công nghệ tài chính', 'banking tech', 'công nghệ ngân hàng',
        'insurtech', 'công nghệ bảo hiểm', 'regtech', 'công nghệ quản lý',
        'clean tech', 'công nghệ sạch', 'green tech', 'công nghệ xanh',
        'energy tech', 'công nghệ năng lượng', 'renewable energy', 'năng lượng tái tạo',
        # THÊM TỪ KHÓA MỚI
        'thực tế ảo', 'vr', 'virtual reality', 'công nghệ thực tế ảo',
        'thực tế tăng cường', 'ar', 'augmented reality'
    ],

    # Thời sự - Tin tức (45+ từ khóa)
    'thời sự': [
        'thời sự', 'tin tức', 'news', 'bản tin', 'tin mới', 'tin nóng',
        'tin nổi bật', 'sự kiện', 'event', 'hội nghị', 'hội thảo',
        'chính trị', 'chính phủ', 'quốc hội', 'thủ tướng', 'chủ tịch',
        'bộ trưởng', 'lãnh đạo', 'đảng', 'nhà nước', 'chính sách',
        'luật', 'nghị định', 'thông tư', 'văn bản pháp luật', 'cải cách',
        'cải cách hành chính', 'công vụ', 'công chức', 'viên chức',
        'bầu cử', 'bỏ phiếu', 'ứng cử', 'tranh cử', 'dân chủ',
        'nhân quyền', 'quyền con người', 'tự do', 'dân sinh', 'an sinh',
        'phúc lợi', 'bảo hiểm', 'y tế', 'giáo dục', 'văn hóa'
    ],

    'thời sự vnexpress': [
        'thời sự trong nước', 'thời sự quốc tế', 'chính trị việt nam',
        'đối ngoại', 'ngoại giao', 'quan hệ quốc tế', 'biển đông',
        'chủ quyền', 'lãnh thổ', 'biên giới', 'hải đảo', 'quần đảo',
        'hoàng sa', 'trường sa', 'kinh tế', 'xã hội', 'văn hóa',
        'giáo dục', 'y tế', 'môi trường', 'biến đổi khí hậu',
        'thiên tai', 'bão lũ', 'hạn hán', 'xâm nhập mặn', 'ô nhiễm',
        'an ninh', 'trật tự', 'pháp luật', 'tòa án', 'viện kiểm sát',
        'công an', 'quân đội', 'quốc phòng', 'an ninh mạng', 'tội phạm',
        'ma túy', 'buôn người', 'tham nhũng', 'tiêu cực', 'khiếu nại'
    ],

    # Giáo dục - Đào tạo (45+ từ khóa)
    'giáo dục': [
        'giáo dục', 'đào tạo', 'education', 'training', 'học tập',
        'giảng dạy', 'dạy học', 'giáo viên', 'giáo viên', 'thầy cô',
        'học sinh', 'sinh viên', 'học viên', 'trường học', 'trường lớp',
        'đại học', 'cao đẳng', 'trung học', 'tiểu học', 'mầm non',
        'mẫu giáo', 'nhà trẻ', 'trường công', 'trường tư', 'trường quốc tế',
        'học phí', 'học bổng', 'thi cử', 'kiểm tra', 'thi tốt nghiệp',
        'thi đại học', 'tuyển sinh', 'xét tuyển', 'hồ sơ', 'nguyện vọng',
        'điểm chuẩn', 'điểm thi', 'kết quả', 'bằng cấp', 'chứng chỉ',
        'văn bằng', 'bằng đại học', 'bằng cao đẳng', 'bằng thạc sĩ',
        'bằng tiến sĩ', 'luận văn', 'luận án', 'nghiên cứu sinh'
    ],

    'giáo dục vnexpress': [
        'tuyển sinh đại học', 'tuyển sinh cao đẳng', 'thi tốt nghiệp thpt',
        'thi đánh giá năng lực', 'thi đánh giá tư duy', 'xét học bạ',
        'nguyện vọng 1', 'nguyện vọng 2', 'nguyện vọng 3', 'điểm sàn',
        'điểm chuẩn', 'trúng tuyển', 'nhập học', 'học bổng', 'học phí',
        'chính sách hỗ trợ', 'sinh viên', 'giảng viên', 'giáo sư',
        'phó giáo sư', 'tiến sĩ', 'thạc sĩ', 'cử nhân', 'kỹ sư',
        'bác sĩ', 'dược sĩ', 'kiến trúc sư', 'luật sư', 'công nghệ thông tin',
        'khoa học máy tính', 'điện tử viễn thông', 'cơ khí', 'xây dựng',
        'kinh tế', 'tài chính ngân hàng', 'quản trị kinh doanh', 'marketing'
    ],

    # Kinh doanh - Tài chính (45+ từ khóa)
    'kinh doanh': [
        'kinh doanh', 'business', 'doanh nghiệp', 'công ty', 'tập đoàn',
        'công ty cổ phần', 'công ty tnhh', 'công ty liên doanh',
        'công ty nước ngoài', 'công ty việt nam', 'startup', 'khởi nghiệp',
        'entrepreneur', 'nhà đầu tư', 'investor', 'quỹ đầu tư', 'venture capital',
        'vốn', 'capital', 'funding', 'tài trợ', 'huy động vốn', 'ipo',
        'cổ phiếu', 'chứng khoán', 'stock', 'thị trường chứng khoán',
        'chỉ số', 'index', 'vn-index', 'hose', 'hnx', 'upcom',
        'trái phiếu', 'bond', 'lãi suất', 'interest rate', 'tỷ giá',
        'exchange rate', 'usd', 'eur', 'jpy', 'cny', 'vàng', 'gold'
    ],

    'kinh doanh vnexpress': [
        'thị trường', 'thương mại', 'xuất khẩu', 'nhập khẩu', 'xuất nhập khẩu',
        'thương mại quốc tế', 'fdi', 'đầu tư nước ngoài', 'oda', 'viện trợ',
        'hợp tác quốc tế', 'hiệp định thương mại', 'fta', 'cptpp', 'rcep',
        'evfta', 'ukfta', 'thuế quan', 'tariff', 'hạn ngạch', 'quota',
        'bảo hộ', 'bảo hộ mậu dịch', 'chống bán phá giá', 'chống trợ cấp',
        'kiện', 'tranh chấp thương mại', 'wto', 'tổ chức thương mại thế giới',
        'asean', 'eu', 'nafta', 'usmca', 'mercosur', 'sàn giao dịch',
        'sàn chứng khoán', 'broker', 'môi giới', 'nhà môi giới', 'trader'
    ],

    # Thể thao (45+ từ khóa)
    'thể thao': [
        'thể thao', 'sports', 'bóng đá', 'football', 'soccer', 'premier league',
        'la liga', 'serie a', 'bundesliga', 'ligue 1', 'v-league',
        'world cup', 'euro', 'asian cup', 'afc', 'fifa', 'uefa',
        'cầu thủ', 'player', 'huấn luyện viên', 'coach', 'đội tuyển',
        'đội bóng', 'câu lạc bộ', 'club', 'trận đấu', 'match',
        'giải đấu', 'tournament', 'champions league', 'europa league',
        'bóng rổ', 'basketball', 'nba', 'bóng chuyền', 'volleyball',
        'cầu lông', 'badminton', 'tennis', 'quần vợt', 'golf',
        'đua xe', 'racing', 'formula 1', 'f1', 'moto gp', 'điền kinh'
    ],

    'thể thao vnexpress': [
        'sea games', 'asian games', 'olympic', 'paralympic', 'thế vận hội',
        'đại hội thể thao', 'huy chương', 'huy chương vàng', 'huy chương bạc',
        'huy chương đồng', 'kỷ lục', 'kỷ lục thế giới', 'kỷ lục châu á',
        'kỷ lục đông nam á', 'kỷ lục việt nam', 'vận động viên', 'athlete',
        'tuyển thủ', 'đội tuyển quốc gia', 'u23', 'u19', 'u16',
        'trẻ', 'trẻ em', 'thanh niên', 'thiếu niên', 'nhi đồng',
        'thể thao học đường', 'thể thao trường học', 'thể thao đại học',
        'thể thao chuyên nghiệp', 'thể thao nghiệp dư', 'thể thao quần chúng',
        'thể thao cho mọi người', 'thể thao phong trào', 'thể thao thành tích cao'
    ],

    # Giải trí (45+ từ khóa)
    'giải trí': [
        'giải trí', 'entertainment', 'phim', 'movie', 'film', 'cinema',
        'điện ảnh', 'hollywood', 'bollywood', 'phim việt', 'phim hàn',
        'phim mỹ', 'phim trung', 'phim nhật', 'phim thái', 'phim ấn độ',
        'diễn viên', 'actor', 'actress', 'đạo diễn', 'director',
        'nhà sản xuất', 'producer', 'biên kịch', 'screenwriter', 'kịch bản',
        'âm nhạc', 'music', 'nhạc', 'bài hát', 'song', 'ca sĩ', 'singer',
        'nhạc sĩ', 'composer', 'nhóm nhạc', 'band', 'concert', 'live show',
        'sân khấu', 'stage', 'kịch', 'drama', 'hài kịch', 'comedy'
    ],

    'giải trí vnexpress': [
        'sao', 'ngôi sao', 'celebrity', 'người nổi tiếng', 'famous',
        'nghệ sĩ', 'artist', 'ca sĩ việt nam', 'diễn viên việt nam',
        'đạo diễn việt nam', 'nhà sản xuất việt nam', 'biên kịch việt nam',
        'phim truyền hình', 'tv series', 'phim lẻ', 'phim bộ', 'phim ngắn',
        'phim tài liệu', 'documentary', 'phim hoạt hình', 'animation',
        'phim kinh dị', 'horror', 'phim hành động', 'action', 'phim tình cảm',
        'romance', 'phim hài', 'comedy', 'phim khoa học viễn tưởng', 'sci-fi',
        'phim viễn tưởng', 'fantasy', 'phim phiêu lưu', 'adventure'
    ],

    # VTV24 (45+ từ khóa)
    'VTV24': [
        'vtv24', 'vtv', 'đài truyền hình việt nam', 'truyền hình quốc hội',
        'thời sự quốc hội', 'phiên họp quốc hội', 'chính phủ', 'thủ tướng',
        'bộ trưởng', 'lãnh đạo', 'đối ngoại', 'ngoại giao', 'quan hệ quốc tế',
        'biển đông', 'chủ quyền', 'lãnh thổ', 'kinh tế', 'xã hội', 'văn hóa',
        'giáo dục', 'y tế', 'môi trường', 'an ninh', 'trật tự', 'pháp luật',
        'tòa án', 'công an', 'quân đội', 'quốc phòng', 'an ninh mạng',
        'tội phạm', 'ma túy', 'thuốc lá', 'rượu bia', 'cờ bạc', 'mại dâm',
        'tham nhũng', 'tiêu cực', 'khiếu nại', 'tố cáo', 'khiếu kiện'
    ],

    # Tin Tức Việt (45+ từ khóa)
    'TIN TỨC VIỆT': [
        'tin việt nam', 'tin trong nước', 'tin địa phương', 'hà nội',
        'hồ chí minh', 'đà nẵng', 'hải phòng', 'cần thơ', 'an giang',
        'bà rịa vũng tàu', 'bắc giang', 'bắc kạn', 'bạc liêu', 'bắc ninh',
        'bến tre', 'bình định', 'bình dương', 'bình phước', 'bình thuận',
        'cà mau', 'cao bằng', 'đắk lắk', 'đắk nông', 'điện biên', 'đồng nai',
        'đồng tháp', 'gia lai', 'hà giang', 'hà nam', 'hà tĩnh', 'hải dương',
        'hậu giang', 'hòa bình', 'hưng yên', 'khánh hòa', 'kiên giang',
        'kon tum', 'lai châu', 'lâm đồng', 'lạng sơn', 'lào cai', 'long an'
    ],

    # THÊM CHỦ ĐỀ MỚI - THỰC TẾ ẢO
    'thực tế ảo': [
        'thực tế ảo', 'vr', 'virtual reality', 'công nghệ thực tế ảo',
        'thực tế ảo vr', 'kính thực tế ảo', 'thiết bị thực tế ảo',
        'game thực tế ảo', 'ứng dụng thực tế ảo', 'phát triển thực tế ảo',
        'công nghệ vr', 'virtual reality technology', 'vr headset',
        'meta quest', 'oculus', 'htc vive', 'playstation vr',
        'thực tế ảo trong giáo dục', 'thực tế ảo y tế', 'thực tế ảo du lịch',
        'thực tế ảo bất động sản', 'thực tế ảo quân sự', 'thực tế ảo thể thao',
        'metaverse', 'vũ trụ ảo', 'web3', 'nft', 'blockchain',
        'augmented reality', 'ar', 'thực tế tăng cường', 'mixed reality',
        'mr', 'thực tế hỗn hợp', 'extended reality', 'xr'
    ],

    # THÊM CHỦ ĐỀ AI VÀ TRÍ TUỆ NHÂN TẠO
    'ai': [
        'ai', 'trí tuệ nhân tạo', 'artificial intelligence', 'machine learning',
        'học máy', 'deep learning', 'học sâu', 'neural network', 'mạng nơ-ron',
        'chatgpt', 'openai', 'gpt', 'llm', 'large language model',
        'computer vision', 'thị giác máy tính', 'natural language processing',
        'xử lý ngôn ngữ tự nhiên', 'nlp', 'robotics', 'người máy',
        'autonomous vehicles', 'xe tự lái', 'smart home', 'nhà thông minh',
        'iot', 'internet of things', 'internet vạn vật', 'big data',
        'dữ liệu lớn', 'data science', 'khoa học dữ liệu', 'data analytics',
        'phân tích dữ liệu', 'predictive analytics', 'phân tích dự đoán'
    ]
}

# ==============================================
# LỚP XỬ LÝ ẢNH VÀ VIDEO NÂNG CAO
# ==============================================

class AdvancedMediaProcessor:
    def __init__(self):
        self.temp_dir = "/tmp/video_processing"
        os.makedirs(self.temp_dir, exist_ok=True)

    def extract_image_from_rss(self, entry):
        """Trích xuất ảnh từ RSS entry với nhiều phương pháp"""
        try:
            print(f"🖼️ Đang trích xuất ảnh từ RSS entry...")

            # PHƯƠNG PHÁP 1: Tìm trong media_content (ưu tiên cao)
            if hasattr(entry, 'media_content') and entry.media_content:
                for media in entry.media_content:
                    if hasattr(media, 'url') and any(ext in media.url.lower() for ext in ['.jpg', '.jpeg', '.png', '.webp']):
                        print(f"✅ Tìm thấy ảnh từ media_content: {media.url}")
                        return media.url

            # PHƯƠNG PHÁP 2: Tìm trong media_thumbnail
            if hasattr(entry, 'media_thumbnail') and entry.media_thumbnail:
                for thumb in entry.media_thumbnail:
                    if hasattr(thumb, 'url') and thumb.url:
                        print(f"✅ Tìm thấy ảnh từ media_thumbnail: {thumb.url}")
                        return thumb.url

            # PHƯƠNG PHÁP 3: Tìm trong links
            if hasattr(entry, 'links'):
                for link in entry.links:
                    # Kiểm tra link có phải là ảnh không
                    if (hasattr(link, 'type') and 'image' in link.type and
                        hasattr(link, 'href') and link.href):
                        print(f"✅ Tìm thấy ảnh từ links: {link.href}")
                        return link.href

            # PHƯƠNG PHÁP 4: Phân tích HTML trong description để tìm ảnh
            description_content = ""
            if hasattr(entry, 'content'):
                for content in entry.content:
                    if hasattr(content, 'value'):
                        description_content += content.value + " "

            if hasattr(entry, 'description'):
                description_content += entry.description

            if hasattr(entry, 'summary'):
                description_content += entry.summary

            if description_content:
                soup = BeautifulSoup(description_content, 'html.parser')
                img_tags = soup.find_all('img')

                for img in img_tags:
                    src = img.get('src')
                    if src and any(ext in src.lower() for ext in ['.jpg', '.jpeg', '.png', '.webp', '//']):
                        # Xử lý URL tương đối
                        if src.startswith('//'):
                            src = 'https:' + src
                        elif src.startswith('/'):
                            # Cần base URL, tạm thời bỏ qua
                            continue

                        print(f"✅ Tìm thấy ảnh từ HTML: {src}")
                        return src

            # PHƯƠNG PHÁP 5: Tìm trong enclosure
            if hasattr(entry, 'enclosures'):
                for enclosure in entry.enclosures:
                    if (hasattr(enclosure, 'type') and 'image' in enclosure.type and
                        hasattr(enclosure, 'href') and enclosure.href):
                        print(f"✅ Tìm thấy ảnh từ enclosure: {enclosure.href}")
                        return enclosure.href

            print("❌ Không tìm thấy ảnh trong RSS entry")
            return None

        except Exception as e:
            print(f"❌ Lỗi trích xuất ảnh từ RSS: {e}")
            return None

    def extract_frame_from_video(self, video_url: str, timestamp_seconds: int = 5):
        """Trích xuất khung hình từ video"""
        try:
            print(f"🎬 Đang trích xuất frame từ: {video_url}")

            video_id = str(uuid.uuid4())[:8]
            temp_video_path = os.path.join(self.temp_dir, f"temp_{video_id}.mp4")
            output_image_path = os.path.join(self.temp_dir, f"frame_{video_id}.jpg")

            # Tải video bằng yt-dlp với timeout
            ydl_opts = {
                'outtmpl': temp_video_path,
                'format': 'best[height<=720]',
                'quiet': True,
                'socket_timeout': 30,
                'retries': 3
            }

            try:
                with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                    ydl.download([video_url])
            except Exception as e:
                print(f"⚠️ Không thể tải video: {e}")
                return None

            # Trích xuất frame
            try:
                cap = cv2.VideoCapture(temp_video_path)
                if not cap.isOpened():
                    return None

                fps = cap.get(cv2.CAP_PROP_FPS)
                if fps <= 0:
                    fps = 30  # Default FPS

                target_frame = int(fps * timestamp_seconds)

                cap.set(cv2.CAP_PROP_POS_FRAMES, target_frame)
                success, frame = cap.read()
                cap.release()

                if success:
                    cv2.imwrite(output_image_path, frame, [cv2.IMWRITE_JPEG_QUALITY, 85])
                    print(f"✅ Đã trích xuất frame: {output_image_path}")

                    # Dọn dẹp
                    if os.path.exists(temp_video_path):
                        os.remove(temp_video_path)

                    return output_image_path
                return None
            except Exception as e:
                print(f"❌ Lỗi trích xuất frame: {e}")
                return None

        except Exception as e:
            print(f"❌ Lỗi trích xuất frame từ video: {e}")
            return None

    def download_youtube_video(self, video_url: str):
        """Tải video YouTube"""
        try:
            print(f"📥 Đang tải video: {video_url}")
            yt = YouTube(video_url)
            stream = yt.streams.filter(progressive=True, file_extension='mp4').order_by('resolution').desc().first()

            if stream:
                video_path = stream.download(output_path=self.temp_dir)
                return video_path
            return None
        except Exception as e:
            print(f"❌ Lỗi tải video: {e}")
            return None

    def create_fallback_image(self, title, category):
        """Tạo ảnh fallback với gradient và text"""
        try:
            # Tạo ảnh gradient với kích thước 600x400
            width, height = 600, 400

            # Màu gradient theo category
            color_map = {
                'thời sự': [(102, 126, 234), (118, 75, 162)],  # Blue to Purple
                'kinh doanh': [(255, 153, 102), (255, 94, 98)],  # Orange to Red
                'công nghệ': [(102, 234, 195), (118, 162, 75)],  # Green to Light Green
                'thể thao': [(255, 204, 102), (255, 153, 51)],  # Yellow to Orange
                'giải trí': [(204, 153, 255), (153, 102, 255)],  # Light Purple to Purple
                'giáo dục': [(102, 217, 255), (0, 140, 186)],   # Light Blue to Blue
            }

            colors = color_map.get(category, [(102, 126, 234), (118, 75, 162)])

            # Tạo gradient
            image = np.zeros((height, width, 3), dtype=np.uint8)
            for i in range(width):
                ratio = i / width
                r = int(colors[0][0] * (1 - ratio) + colors[1][0] * ratio)
                g = int(colors[0][1] * (1 - ratio) + colors[1][1] * ratio)
                b = int(colors[0][2] * (1 - ratio) + colors[1][2] * ratio)
                image[:, i] = [b, g, r]

            # Thêm text
            font = cv2.FONT_HERSHEY_SIMPLEX
            text = title[:50] + "..." if len(title) > 50 else title

            # Tính toán vị trí text
            text_size = cv2.getTextSize(text, font, 0.8, 2)[0]
            text_x = (width - text_size[0]) // 2
            text_y = (height + text_size[1]) // 2

            # Thêm shadow
            cv2.putText(image, text, (text_x+2, text_y+2), font, 0.8, (0, 0, 0), 2, cv2.LINE_AA)
            # Thêm text chính
            cv2.putText(image, text, (text_x, text_y), font, 0.8, (255, 255, 255), 2, cv2.LINE_AA)

            # Thêm icon category
            icon_text = CATEGORY_ICONS.get(category, '📰')
            icon_size = cv2.getTextSize(icon_text, font, 1.5, 3)[0]
            icon_x = (width - icon_size[0]) // 2
            icon_y = text_y - 50

            cv2.putText(image, icon_text, (icon_x, icon_y), font, 1.5, (255, 255, 255), 3, cv2.LINE_AA)

            # Lưu ảnh
            image_path = os.path.join(self.temp_dir, f"fallback_{uuid.uuid4().hex[:8]}.jpg")
            cv2.imwrite(image_path, image)

            print(f"✅ Đã tạo ảnh fallback: {image_path}")
            return image_path

        except Exception as e:
            print(f"❌ Lỗi tạo ảnh fallback: {e}")
            return None

    def cleanup_temp_files(self):
        """Dọn dẹp file tạm"""
        try:
            for file in os.listdir(self.temp_dir):
                file_path = os.path.join(self.temp_dir, file)
                if os.path.isfile(file_path):
                    os.remove(file_path)
            print("✅ Đã dọn dẹp file tạm")
        except Exception as e:
            print(f"❌ Lỗi dọn dẹp: {e}")

# ==============================================
# LỚP XỬ LÝ NỘI DUNG THÔNG MINH - ĐÃ SỬA LỖI VTV24 VÀ THÊM PHƯƠNG THỨC THIẾU
# ==============================================

class ContentGenerator:
    def __init__(self):
        self.content_templates = self.initialize_templates()

    def initialize_templates(self):
        """Khởi tạo template nội dung"""
        return {
            'thời sự': {
                'intro': [
                    "Sự kiện {title} đang thu hút sự quan tâm đặc biệt của công chúng với những diễn biến mới nhất.",
                    "Theo thông tin mới nhất, {title} đang diễn biến phức tạp và được dư luận quan tâm.",
                    "Vấn đề {title} đang được dư luận đặc biệt quan tâm trong những ngày qua."
                ]
            },
            'kinh doanh': {
                'intro': [
                    "Thị trường đang có phản ứng tích cực trước thông tin về {title}.",
                    "Sự kiện {title} đang tác động mạnh mẽ đến hoạt động kinh doanh và đầu tư.",
                    "Theo các nguồn tin, {title} đang ảnh hưởng sâu rộng đến thị trường tài chính."
                ]
            },
            'công nghệ': {
                'intro': [
                    "Công nghệ {title} đang mở ra những hướng phát triển mới đầy hứa hẹn.",
                    "Sự kiện {title} trong ngành công nghệ đang được giới chuyên môn đánh giá cao.",
                    "Theo thông tin mới, {title} đang định hình xu hướng công nghệ tương lai."
                ]
            },
            'thể thao': {
                'intro': [
                    "Sự kiện thể thao {title} đang thu hút sự chú ý đặc biệt của người hâm mộ.",
                    "Theo thông tin mới nhất, {title} đang diễn ra sôi động với nhiều bất ngờ.",
                    "Giới thể thao đang quan tâm sâu sắc đến sự kiện {title} với nhiều kỳ vọng."
                ]
            },
            'giáo dục': {
                'intro': [
                    "Vấn đề giáo dục {title} đang nhận được sự quan tâm đặc biệt từ phụ huynh và học sinh.",
                    "Theo thông tin mới nhất, {title} trong lĩnh vực giáo dục đang có những diễn biến quan trọng.",
                    "Sự kiện {title} đang tác động mạnh mẽ đến hệ thống giáo dục và đào tạo."
                ]
            }
        }

    def clean_title_for_content(self, title):
        """Làm sạch tiêu đề cho nội dung bài báo - LOẠI BỎ VTV24"""
        # Loại bỏ các ký hiệu không cần thiết và chuỗi VTV24
        cleaned = re.sub(r'\s*\|\s*VTV24\s*', '', title, flags=re.IGNORECASE)
        cleaned = re.sub(r'\s*VTV24\s*', '', cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r'\s*\|\s*', ' - ', cleaned)
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        return cleaned

    def generate_article_content(self, title, base_content, category, min_words=600):
        """Tạo nội dung bài báo chi tiết với 5 góc nhìn khác biệt"""
        try:
            # Làm sạch tiêu đề và nội dung gốc
            clean_title = self.clean_title_for_content(title)
            clean_content = self.clean_text(base_content)

            # Tạo bài báo theo cấu trúc báo chí chuyên nghiệp
            article_content = self.create_journalistic_article(clean_title, clean_content, category, min_words)

            return article_content

        except Exception as e:
            print(f"❌ Lỗi tạo nội dung: {e}")
            return self.generate_comprehensive_article(title, category, min_words)

    def create_journalistic_article(self, title, content, category, min_words):
        """Tạo bài báo theo phong cách báo chí với cấu trúc cụ thể"""
        try:
            # Tạo các phần của bài báo
            intro = self.generate_journalistic_intro(title, category)
            main_content = self.generate_detailed_content(title, content, category, min_words - 200)
            conclusion = self.generate_journalistic_conclusion(title, category)

            # Kết hợp thành bài báo hoàn chỉnh
            full_content = f"{intro}\n\n{main_content}\n\n{conclusion}"

            # Đảm bảo đủ số từ
            current_words = len(full_content.split())
            if current_words < min_words:
                additional_content = self.generate_additional_journalistic_content(title, category, min_words - current_words)
                full_content += f"\n\n{additional_content}"

            return full_content

        except Exception as e:
            print(f"❌ Lỗi tạo bài báo báo chí: {e}")
            return self.generate_fallback_journalistic_article(title, category, min_words)

    def generate_journalistic_intro(self, title, category):
        """Tạo phần mở đầu theo phong cách báo chí"""
        locations = {
            'thời sự': ['Hà Nội', 'TP.HCM', 'Đà Nẵng', 'Hải Phòng', 'Cần Thơ'],
            'kinh doanh': ['thị trường chứng khoán', 'thị trường bất động sản', 'ngành ngân hàng', 'thị trường tiêu dùng'],
            'công nghệ': ['thung lũng Silicon', 'các startup công nghệ', 'ngành công nghiệp 4.0', 'cộng đồng công nghệ'],
            'thể thao': ['sân vận động', 'giải đấu', 'câu lạc bộ', 'trung tâm thể thao'],
            'giáo dục': ['các trường đại học', 'Bộ Giáo dục và Đào tạo', 'các cơ sở giáo dục', 'các trường học']
        }

        location = random.choice(locations.get(category, locations['thời sự']))
        current_date = self.get_current_time().split()[0]  # Lấy ngày tháng

        intros = [
            f"{location} - {current_date}, {title} đang thu hút sự quan tâm đặc biệt của dư luận với những diễn biến mới nhất. Tình hình này đang tác động trực tiếp đến đời sống người dân và được các chuyên gia đánh giá có ý nghĩa quan trọng.",
            f"{current_date} - Sự kiện {title} tại {location} đang diễn biến phức tạp, buộc các cơ quan chức năng phải vào cuộc xử lý. Thông tin mới nhất cho thấy vấn đề này sẽ còn tiếp tục ảnh hưởng trong thời gian tới.",
            f"{location} - Trước tình hình {title}, người dân đang tỏ ra lo lắng và kêu gọi sự vào cuộc kịp thời của chính quyền. {current_date}, các chuyên gia đã có những đánh giá ban đầu về sự việc."
        ]

        return random.choice(intros)

    def generate_detailed_content(self, title, base_content, category, target_words):
        """Tạo nội dung chi tiết với 5 góc nhìn khác biệt"""
        try:
            paragraphs = []
            current_words = 0

            # GÓC NHÌN 1: Phân tích chuyên sâu từ góc độ chuyên môn
            expert_analysis = self.generate_expert_analysis(title, category)
            paragraphs.append(f"### 🔍 Góc Nhìn Chuyên Gia\n\n{expert_analysis}")
            current_words += len(expert_analysis.split())

            # GÓC NHÌN 2: Đánh giá tác động thực tiễn
            practical_impact = self.generate_practical_impact(title, category)
            paragraphs.append(f"### 📊 Góc Nhìn Thực Tiễn\n\n{practical_impact}")
            current_words += len(practical_impact.split())

            # GÓC NHÌN 3: Phân tích xu hướng và dự báo
            trend_analysis = self.generate_trend_analysis(title, category)
            paragraphs.append(f"### 📈 Góc Nhìn Xu Hướng\n\n{trend_analysis}")
            current_words += len(trend_analysis.split())

            # GÓC NHÌN 4: Đánh giá từ góc độ cộng đồng
            community_perspective = self.generate_community_perspective(title, category)
            paragraphs.append(f"### 👥 Góc Nhìn Cộng Đồng\n\n{community_perspective}")
            current_words += len(community_perspective.split())

            # GÓC NHÌN 5: Phân tích giải pháp và khuyến nghị
            solution_perspective = self.generate_solution_perspective(title, category)
            paragraphs.append(f"### 💡 Góc Nhìn Giải Pháp\n\n{solution_perspective}")
            current_words += len(solution_perspective.split())

            # Thêm các đoạn bổ sung nếu cần đủ số từ
            while current_words < target_words:
                additional_para = self.generate_additional_perspective(title, category)
                paragraphs.append(additional_para)
                current_words += len(additional_para.split())

            return "\n\n".join(paragraphs)

        except Exception as e:
            print(f"❌ Lỗi tạo nội dung chi tiết: {e}")
            return self.generate_fallback_content(title, category, target_words)

    def generate_expert_analysis(self, title, category):
        """Góc nhìn chuyên gia - phân tích chuyên sâu"""
        expert_templates = {
            'thời sự': [
                f"Theo phân tích chuyên sâu từ các chuyên gia hàng đầu, {title} cho thấy những điểm đáng chú ý về mặt chính sách và quản lý. Các chuyên gia nhấn mạnh đây không chỉ là vấn đề đơn lẻ mà phản ánh những thách thức lớn hơn trong công tác quản lý đô thị và an sinh xã hội. Phân tích dữ liệu cho thấy xu hướng gia tăng cần được quan tâm đặc biệt từ các cơ quan chức năng. Các giải pháp cần được xây dựng trên cơ sở khoa học và thực tiễn, kết hợp kinh nghiệm quốc tế và đặc thù địa phương.",
                f"Góc nhìn chuyên môn về {title} chỉ ra những nguyên nhân sâu xa liên quan đến cơ chế phối hợp giữa các ban ngành. Các chuyên gia cho rằng cần có sự điều chỉnh trong cách tiếp cận giải quyết vấn đề, từ phản ứng sang chủ động phòng ngừa. Nghiên cứu thực tiễn cho thấy hiệu quả của các biện pháp can thiệp phụ thuộc lớn vào tính đồng bộ và sự tham gia của cộng đồng. Các bài học kinh nghiệm từ mô hình thành công cần được nhân rộng và điều chỉnh phù hợp."
            ],
            'kinh doanh': [
                f"Phân tích chuyên sâu từ các chuyên gia kinh tế về {title} cho thấy tác động đa chiều đến cấu trúc thị trường. Các chỉ số kinh tế vĩ mô phản ánh mức độ ảnh hưởng sâu rộng của sự kiện này đến chuỗi cung ứng và thị trường tài chính. Chuyên gia nhấn mạnh nhu cầu cấp thiết về việc điều chỉnh chiến lược kinh doanh trong bối cảnh mới. Phân tích xu hướng dài hạn chỉ ra những cơ hội và thách thức mà doanh nghiệp cần lưu ý để phát triển bền vững.",
                f"Góc nhìn chuyên môn về {title} tập trung vào phân tích tác động đến cạnh tranh thị trường và hành vi người tiêu dùng. Các chuyên gia chỉ ra sự thay đổi trong mô hình kinh doanh và sự dịch chuyển trong cấu trúc ngành. Nghiên cứu cho thấy tầm quan trọng của việc ứng dụng công nghệ và đổi mới sáng tạo trong việc thích ứng với biến động. Các khuyến nghị chính sách cần dựa trên bằng chứng thực tiễn và kinh nghiệm quốc tế."
            ],
            'giáo dục': [
                f"Phân tích từ các chuyên gia giáo dục về {title} cho thấy những tác động sâu sắc đến hệ thống đào tạo và chất lượng giáo dục. Các chuyên gia nhấn mạnh sự cần thiết của việc đổi mới phương pháp giảng dạy và cập nhật chương trình đào tạo phù hợp với xu hướng phát triển. Nghiên cứu cho thấy sự thay đổi trong nhu cầu của thị trường lao động đòi hỏi các cơ sở giáo dục phải điều chỉnh để đáp ứng.",
                f"Góc nhìn chuyên môn về {title} tập trung vào phân tích các yếu tố ảnh hưởng đến chất lượng đào tạo và cơ hội phát triển của học sinh, sinh viên. Các chuyên gia chỉ ra sự cần thiết của việc kết hợp giữa lý thuyết và thực hành, cũng như việc ứng dụng công nghệ trong giáo dục. Các giải pháp cần được xây dựng dựa trên nghiên cứu khoa học và thực tiễn giáo dục."
            ]
        }
        templates = expert_templates.get(category, expert_templates['thời sự'])
        return random.choice(templates)

    def generate_practical_impact(self, title, category):
        """Góc nhìn thực tiễn - tác động cụ thể"""
        practical_templates = {
            'thời sự': [
                f"Trên thực tế, {title} đã tác động trực tiếp đến đời sống hàng ngày của người dân tại nhiều khu vực. Các hoạt động sản xuất, kinh doanh và sinh hoạt chịu ảnh hưởng không nhỏ, đòi hỏi sự điều chỉnh linh hoạt từ cộng đồng. Ghi nhận thực địa cho thấy mức độ nghiêm trọng của vấn đề và những khó khăn thực tế mà người dân đang phải đối mặt. Các giải pháp tình thế đã được áp dụng nhưng hiệu quả vẫn cần được đánh giá toàn diện hơn.",
                f"Tác động thực tiễn của {title} thể hiện rõ qua những con số thống kê và phản ánh từ cơ sở. Các dịch vụ công thiết yếu chịu áp lực lớn, trong khi nhu cầu của người dân tiếp tục gia tăng. Thực tế cho thấy sự cần thiết của các biện pháp hỗ trợ kịp thời và hiệu quả để ổn định đời sống. Các bài học kinh nghiệm từ thực tiễn cần được tổng kết và áp dụng rộng rãi."
            ],
            'kinh doanh': [
                f"Tác động thực tiễn của {title} đến hoạt động sản xuất kinh doanh được thể hiện qua các chỉ số tài chính cụ thể. Doanh thu, lợi nhuận và thị phần của nhiều doanh nghiệp chịu ảnh hưởng đáng kể, đòi hỏi các điều chỉnh chiến lược kịp thời. Thực tế thị trường cho thấy sự thay đổi trong hành vi tiêu dùng và cấu trúc cạnh tranh ngành. Các doanh nghiệp buộc phải thích ứng nhanh để duy trì hoạt động và phát triển.",
                f"Trên thực tế, {title} đã tạo ra những thay đổi cụ thể trong môi trường kinh doanh và đầu tư. Các dự án đầu tư chịu tác động trực tiếp, trong khi kế hoạch sản xuất kinh doanh cần điều chỉnh linh hoạt. Thống kê cho thấy mức độ ảnh hưởng đến việc làm và thu nhập của người lao động. Các biện pháp hỗ trợ thực tế cần được triển khai đồng bộ để giảm thiểu tác động tiêu cực."
            ],
            'giáo dục': [
                f"Trên thực tế, {title} đã tác động trực tiếp đến hàng nghìn học sinh, sinh viên và các cơ sở giáo dục. Các hoạt động giảng dạy, học tập và đánh giá kết quả chịu ảnh hưởng đáng kể, đòi hỏi sự điều chỉnh linh hoạt từ nhà trường và người học. Thực tế cho thấy nhu cầu cấp thiết về việc cải tiến phương pháp giảng dạy và nâng cao chất lượng đào tạo.",
                f"Tác động thực tiễn của {title} thể hiện rõ qua những thay đổi trong hệ thống giáo dục và đào tạo. Các cơ sở giáo dục phải đối mặt với nhiều thách thức mới, trong khi nhu cầu về chất lượng đào tạo ngày càng cao. Thực tế cho thấy sự cần thiết của việc đầu tư vào cơ sở vật chất và nâng cao năng lực đội ngũ giảng viên."
            ]
        }
        templates = practical_templates.get(category, practical_templates['thời sự'])
        return random.choice(templates)

    def generate_trend_analysis(self, title, category):
        """Góc nhìn xu hướng - phân tích dài hạn"""
        trend_templates = {
            'thời sự': [
                f"Phân tích xu hướng cho thấy {title} có khả năng tiếp tục diễn biến phức tạp trong thời gian tới. Các yếu tố khách quan như biến đổi khí hậu và tốc độ đô thị hóa sẽ tiếp tục tác động đến vấn đề này. Dự báo từ các mô hình phân tích chỉ ra khả năng mở rộng phạm vi ảnh hưởng nếu không có biện pháp can thiệp hiệu quả. Xu hướng này đòi hỏi cách tiếp cận tổng thể và dài hạn từ các cơ quan quản lý.",
                f"Góc nhìn xu hướng về {title} chỉ ra sự thay đổi trong mô hình và quy mô của vấn đề theo thời gian. Phân tích dữ liệu lịch sử cho thấy chu kỳ và mức độ gia tăng của sự việc, cung cấp cơ sở cho công tác dự báo và phòng ngừa. Xu hướng phát triển trong tương lai sẽ chịu ảnh hưởng của nhiều yếu tố kinh tế - xã hội và môi trường. Các kịch bản phát triển cần được xây dựng để chủ động ứng phó."
            ],
            'kinh doanh': [
                f"Xu hướng thị trường sau {title} cho thấy sự dịch chuyển mạnh mẽ trong cấu trúc ngành và mô hình kinh doanh. Phân tích dài hạn chỉ ra những thay đổi căn bản trong cách thức vận hành và cạnh tranh trên thị trường. Các xu hướng công nghệ và tiêu dùng mới sẽ tiếp tục định hình lại bức tranh kinh doanh trong tương lai. Doanh nghiệp cần nhận diện và thích ứng với các xu hướng này để duy trì lợi thế cạnh tranh.",
                f"Góc nhìn xu hướng về {title} tập trung vào phân tích sự phát triển dài hạn của thị trường và ngành hàng. Các chỉ số kinh tế vĩ mô và vi mô cho thấy hướng đi của nền kinh tế trong bối cảnh mới. Xu hướng toàn cầu và khu vực sẽ tiếp tục tác động đến môi trường kinh doanh trong nước. Phân tích này cung cấp cơ sở quan trọng cho việc hoạch định chiến lược phát triển."
            ],
            'giáo dục': [
                f"Xu hướng giáo dục sau {title} cho thấy sự thay đổi mạnh mẽ trong phương pháp đào tạo và quản lý giáo dục. Phân tích dài hạn chỉ ra sự dịch chuyển từ giáo dục truyền thống sang các mô hình giáo dục hiện đại, linh hoạt hơn. Các xu hướng công nghệ và nhu cầu thị trường lao động sẽ tiếp tục định hình lại hệ thống giáo dục trong tương lai.",
                f"Góc nhìn xu hướng về {title} tập trung vào phân tích sự phát triển của hệ thống giáo dục trong bối cảnh mới. Các yếu tố như công nghệ, toàn cầu hóa và thay đổi nhân khẩu học sẽ tiếp tục tác động đến giáo dục. Xu hướng này đòi hỏi sự đổi mới và thích ứng từ các cơ sở giáo dục để đáp ứng nhu cầu xã hội."
            ]
        }
        templates = trend_templates.get(category, trend_templates['thời sự'])
        return random.choice(templates)

    def generate_community_perspective(self, title, category):
        """Góc nhìn cộng đồng - phản ánh từ người dân"""
        community_templates = {
            'thời sự': [
                f"Từ góc độ cộng đồng, {title} đã tác động sâu sắc đến đời sống và sinh kế của người dân. Các hộ gia đình phải điều chỉnh thói quen sinh hoạt và phương thức sản xuất để thích ứng với tình hình mới. Phản ánh từ người dân cho thấy mức độ quan tâm và lo ngại về những ảnh hưởng lâu dài. Sự tham gia của cộng đồng trong việc đề xuất và thực hiện giải pháp là yếu tố then chốt cho thành công.",
                f"Cộng đồng đang thể hiện sự quan tâm đặc biệt đến {title} với nhiều ý kiến đóng góp và phản ánh. Các tổ chức xã hội và hội nhóm đã tích cực tham gia vào quá trình hỗ trợ và vận động giải quyết vấn đề. Ghi nhận từ cơ sở cho thấy nhu cầu cấp thiết về thông tin minh bạch và sự hỗ trợ kịp thời từ chính quyền. Sức mạnh cộng đồng đang được phát huy để cùng vượt qua thách thức."
            ],
            'kinh doanh': [
                f"Cộng đồng doanh nghiệp đang có những phản ứng đa dạng trước {title}, từ thận trọng đến chủ động thích ứng. Các hiệp hội ngành nghề đã tích cực phối hợp để tìm kiếm giải pháp chung và hỗ trợ thành viên. Phản ánh từ doanh nghiệp vừa và nhỏ cho thấy những khó khăn cụ thể trong việc thích ứng với biến động. Sự chia sẻ kinh nghiệm và hỗ trợ lẫn nhau trong cộng đồng doanh nghiệp đang phát huy hiệu quả.",
                f"Góc nhìn cộng đồng về {title} thể hiện qua sự tham gia của các bên liên quan trong việc tìm kiếm giải pháp. Người tiêu dùng, doanh nghiệp và các tổ chức xã hội đều có tiếng nói trong việc định hướng phát triển. Các sáng kiến từ cộng đồng đang góp phần quan trọng vào việc tháo gỡ khó khăn và tạo động lực phát triển mới. Sự đồng thuận và hợp tác trong cộng đồng là nền tảng cho sự phục hồi bền vững."
            ],
            'giáo dục': [
                f"Từ góc độ cộng đồng, {title} đã nhận được sự quan tâm đặc biệt từ phụ huynh, học sinh và các nhà giáo dục. Các ý kiến phản ánh cho thấy nhu cầu cấp thiết về việc cải thiện chất lượng giáo dục và tạo cơ hội học tập công bằng cho mọi người. Sự tham gia của cộng đồng trong việc đóng góp ý kiến và giám sát chất lượng giáo dục là rất quan trọng.",
                f"Cộng đồng giáo dục đang có những phản ứng tích cực trước {title}, với nhiều sáng kiến và đề xuất cải tiến. Các trường học, giáo viên và phụ huynh đã tích cực phối hợp để tìm ra các giải pháp phù hợp. Phản ánh từ cơ sở cho thấy nhu cầu về sự hỗ trợ và hướng dẫn từ các cơ quan quản lý giáo dục."
            ]
        }
        templates = community_templates.get(category, community_templates['thời sự'])
        return random.choice(templates)

    def generate_solution_perspective(self, title, category):
        """Góc nhìn giải pháp - đề xuất và khuyến nghị"""
        solution_templates = {
            'thời sự': [
                f"Từ góc độ giải pháp, {title} đòi hỏi cách tiếp cận tổng thể và đồng bộ từ nhiều phía. Các biện pháp ngắn hạn cần được triển khai ngay để ổn định tình hình, trong đó ưu tiên bảo vệ quyền lợi người dân. Về trung và dài hạn, cần xây dựng kế hoạch chiến lược với lộ trình rõ ràng, tập trung vào cải cách thể chế và nâng cao năng lực quản lý. Giải pháp công nghệ và sáng tạo cần được ứng dụng mạnh mẽ để nâng cao hiệu quả quản lý.",
                f"Góc nhìn giải pháp cho {title} nhấn mạnh sự cần thiết của việc kết hợp giữa biện pháp kỹ thuật và quản lý. Các giải pháp cần dựa trên nguyên tắc bền vững, đảm bảo cân bằng giữa phát triển kinh tế và bảo vệ môi trường. Sự tham gia của khu vực tư nhân và các tổ chức xã hội sẽ góp phần quan trọng trong việc triển khai hiệu quả các giải pháp. Cơ chế giám sát và đánh giá cần được thiết lập để đảm bảo tính hiệu quả và bền vững."
            ],
            'kinh doanh': [
                f"Giải pháp cho {title} cần được tiếp cận từ cả góc độ vĩ mô và vi mô. Về phía nhà nước, cần hoàn thiện khung pháp lý và chính sách hỗ trợ doanh nghiệp. Về phía doanh nghiệp, cần chủ động đổi mới mô hình kinh doanh và nâng cao năng lực cạnh tranh. Các giải pháp công nghệ và chuyển đổi số cần được đẩy mạnh để tạo đột phá trong hoạt động sản xuất kinh doanh. Sự phối hợp chặt chẽ giữa các bên sẽ tạo ra sức mạnh tổng hợp để vượt qua thách thức.",
                f"Từ góc độ giải pháp, {title} đòi hỏi sự linh hoạt và sáng tạo trong cách tiếp cận. Các giải pháp tài chính cần được triển khai để hỗ trợ thanh khoản và duy trì hoạt động sản xuất kinh doanh. Về chiến lược, doanh nghiệp cần tái cấu trúc và đa dạng hóa thị trường để giảm thiểu rủi ro. Giải pháp nhân sự và phát triển nguồn nhân lực cũng cần được chú trọng để nâng cao năng lực cạnh tranh trong dài hạn."
            ],
            'giáo dục': [
                f"Giải pháp cho {title} cần được tiếp cận toàn diện từ nhiều phía. Về phía nhà nước, cần hoàn thiện chính sách giáo dục và tăng cường đầu tư cho giáo dục. Về phía nhà trường, cần đổi mới phương pháp giảng dạy và nâng cao chất lượng đào tạo. Các giải pháp công nghệ và chuyển đổi số cần được ứng dụng mạnh mẽ trong giáo dục. Sự phối hợp giữa nhà trường, gia đình và xã hội là yếu tố then chốt cho thành công.",
                f"Từ góc độ giải pháp, {title} đòi hỏi sự đổi mới và sáng tạo trong cách tiếp cận. Các giải pháp cần tập trung vào việc nâng cao chất lượng đào tạo, cải tiến phương pháp giảng dạy và tăng cường cơ sở vật chất. Việc phát triển nguồn nhân lực chất lượng cao và đáp ứng nhu cầu thị trường lao động cũng cần được chú trọng."
            ]
        }
        templates = solution_templates.get(category, solution_templates['thời sự'])
        return random.choice(templates)

    def generate_additional_perspective(self, title, category):
        """Tạo thêm góc nhìn bổ sung nếu cần"""
        additional_templates = {
            'thời sự': [
                f"Bên cạnh những góc nhìn đã nêu, {title} còn cần được xem xét dưới góc độ pháp lý và quy hoạch. Các quy định hiện hành cần được rà soát để đảm bảo phù hợp với thực tiễn phát triển. Công tác quy hoạch đô thị và nông thôn cần được điều chỉnh để ứng phó với những thách thức mới. Sự phối hợp giữa các ngành, các cấp cần được tăng cường để tạo sức mạnh tổng hợp trong giải quyết vấn đề.",
                f"Một góc nhìn khác về {title} là từ khía cạnh văn hóa và xã hội. Sự kiện này có thể tác động đến nếp sống, văn hóa ứng xử và mối quan hệ cộng đồng. Việc bảo tồn và phát huy các giá trị văn hóa truyền thống trong bối cảnh mới cũng cần được quan tâm. Các giải pháp cần hài hòa giữa phát triển kinh tế và bảo tồn bản sắc văn hóa, đảm bảo sự phát triển bền vững cho cộng đồng."
            ],
            'kinh doanh': [
                f"Góc nhìn bổ sung cho {title} là từ phía người tiêu dùng và thị trường lao động. Người tiêu dùng đang thay đổi thói quen mua sắm và yêu cầu cao hơn về chất lượng sản phẩm dịch vụ. Thị trường lao động cũng có những biến động với sự dịch chuyển nghề nghiệp và yêu cầu mới về kỹ năng. Các doanh nghiệp cần chú ý đến những thay đổi này để điều chỉnh chiến lược phù hợp.",
                f"Xem xét {title} từ góc độ công nghệ và đổi mới sáng tạo cho thấy cơ hội chuyển đổi số trong các ngành kinh tế. Ứng dụng công nghệ mới có thể giúp doanh nghiệp tối ưi hóa hoạt động, nâng cao năng suất và cải thiện trải nghiệm khách hàng. Đổi mới sáng tạo không chỉ trong sản phẩm mà còn trong mô hình kinhdoanh sẽ tạo lợi thế cạnh tranh bền vững cho doanh nghiệp."
            ],
            'giáo dục': [
                f"Góc nhìn bổ sung cho {title} là từ khía cạnh công nghệ và đổi mới sáng tạo trong giáo dục. Ứcông nghệ thông tin và truyền thông trong giảng dạy và học tập đang tạo ra những thay đổi tích cực. Các mô hình giáo dục trực tuyến và học tập kết hợp đang trở thành xu hướng phổ biến. Việc đổi mới phương pháp giảng dạy và học tập cần được chú trọng để nâng cao hiệu quả giáo dục.",
                f"Xem xét {title} từ góc độ hội nhập quốc tế cho thấy cơ hội và thách thức đối với giáo dục Việt Nam. Sự giao lưu và hợp tác quốc tế trong giáo dục đang mở ra nhiều cơ hội phát triển. Tuy nhiên, cũng đặt ra yêu cầu về việc nâng cao chất lượng đào tạo và đáp ứng chuẩn mực quốc tế. Các cơ sở giáo dục cần chủ động hội nhập để nâng cao vị thế và chất lượng."
            ]
        }
        templates = additional_templates.get(category, additional_templates['thời sự'])
        return random.choice(templates)

    def generate_journalistic_conclusion(self, title, category):
        """Tạo kết luận theo phong cách báo chí"""
        conclusion_templates = {
            'thời sự': [
                f"Sự kiện {title} một lần nữa cho thấy tầm quan trọng của việc xây dựng hệ thống quản lý hiệu quả và nâng cao ý thức cộng đồng. Bài học kinh nghiệm từ sự việc này cần được ghi nhận và áp dụng để ngăn chặn những tình huống tương tự trong tương lai. Người dân và chính quyền cần tiếp tục phối hợp chặt chẽ để tìm ra giải pháp bền vững.",
                f"Trước mắt, các biện pháp xử lý {title} đang được triển khai quyết liệt. Tuy nhiên, về lâu dài, cần có sự thay đổi căn bản trong cách tiếp cận và giải quyết vấn đề. Sự chung tay của toàn xã hội sẽ tạo ra sức mạnh tổng hợp để vượt qua thách thức và xây dựng cộng đồng ngày càng tốt đẹp hơn.",
                f"{title} không chỉ là vấn đề trước mắt mà còn đặt ra những thách thức lâu dài về quản lý và phát triển. Việc rút kinh nghiệm và hoàn thiện cơ chế, chính sách là hết sức cần thiết. Cộng đồng kỳ vọng vào sự chuyển biến tích cực và những kết quả cụ thể trong thời gian tới."
            ],
            'kinh doanh': [
                f"{title} đang định hình lại bức tranh kinh doanh và đầu tư. Các doanh nghiệp cần nhanh chóng thích ứng với xu hướng mới và tìm kiếm cơ hội trong thách thức. Sự linh hoạt và sáng tạo sẽ là yếu tố then chốt để thành công trong giai đoạn chuyển đổi quan trọng này.",
                f"Bối cảnh kinh doanh sau {title} sẽ có nhiều thay đổi so với trước đây. Các doanh nghiệp cần chuẩn bị cho những kịch bản phát triển mới và xây dựng năng lực cạnh tranh bền vững. Việc nắm bắt cơ hội và quản lý rủi ro hiệu quả sẽ quyết định vị thế của doanh nghiệp trong tương lai.",
                f"Sự kiện {title} cho thấy tính phức tạp và biến động của môi trường kinh doanh hiện đại. Các bài học kinh nghiệm cần được đúc kết và áp dụng vào chiến lược phát triển. Sự chủ động và khả năng thích ứng sẽ giúp doanh nghiệp không chỉ tồn tại mà còn phát triển mạnh mẽ trong bối cảnh mới."
            ],
            'giáo dục': [
                f"{title} đang định hình lại tương lai của giáo dục và đào tạo. Các cơ sở giáo dục cần nhanh chóng thích ứng với xu hướng mới và tìm kiếm cơ hội đổi mới. Sự sáng tạo và đổi mới sẽ là yếu tố then chốt để nâng cao chất lượng giáo dục trong giai đoạn phát triển quan trọng này.",
                f"Bối cảnh giáo dục sau {title} sẽ có nhiều thay đổi so với trước đây. Các cơ sở giáo dục cần chuẩn bị cho những mô hình đào tạo mới và xây dựng năng lực cạnh tranh bền vững. Việc nắm bắt cơ hội và ứng phó với thách thức hiệu quả sẽ quyết định chất lượng giáo dục trong tương lai.",
                f"Sự kiện {title} cho thấy tính cấp thiết của việc đổi mới giáo dục trong bối cảnh hiện đại. Các bài học kinh nghiệm cần được đúc kết và áp dụng vào chiến lược phát triển giáo dục. Sự chủ động và khả năng thích ứng sẽ giúp hệ thống giáo dục không chỉ đáp ứng nhu cầu hiện tại mà còn định hướng cho tương lai."
            ]
        }

        templates = conclusion_templates.get(category, conclusion_templates['thời sự'])
        return random.choice(templates)

    def generate_additional_journalistic_content(self, title, category, word_needed):
        """Tạo nội dung bổ sung để đảm bảo đủ số từ"""
        additional_content = ""
        current_words = 0

        while current_words < word_needed:
            paragraph = self.generate_additional_perspective(title, category)
            additional_content += paragraph + "\n\n"
            current_words += len(paragraph.split())

        return additional_content.strip()

    def generate_fallback_journalistic_article(self, title, category, min_words):
        """Tạo bài báo fallback khi có lỗi"""
        current_date = self.get_current_time().split()[0]

        article = f"""Hà Nội - {current_date}, {title} đang thu hút sự quan tâm đặc biệt của dư luận với những diễn biến mới nhất. Tình hình này đang tác động trực tiếp đến đời sống người dân và được các chuyên gia đánh giá có ý nghĩa quan trọng.

### 🔍 Góc Nhìn Chuyên Gia

Theo phân tích chuyên sâu từ các chuyên gia hàng đầu, {title} cho thấy những điểm đáng chú �ết về mặt chính sách và quản lý. Các chuyên gia nhấn mạnh đây không chỉ là vấn đề đơn lẻ mà phản ánh những thách thức lớn hơn trong công tác quản lý đô thị và an sinh xã hội.

### 📊 Góc Nhìn Thực Tiễn

Trên thực tế, {title} đã tác động trực tiếp đến đời sống hàng ngày của người dân tại nhiều khu vực. Các hoạt động sản xuất, kinh doanh và sinh hoạt chịu ảnh hưởng không nhỏ, đòi hỏi sự điều chỉnh linh hoạt từ cộng đồng.

### 📈 Góc Nhìn Xu Hướng

Phân tích xu hướng cho thấy {title} có khả năng tiếp tục diễn biến phức tạp trong thời gian tới. Các yếu tố khách quan như biến đổi khí hậu và tốc độ đô thị hóa sẽ tiếp tục tác động đến vấn đề này.

### 👥 Góc Nhìn Cộng Đồng

Từ góc độ cộng đồng, {title} đã tác động sâu sắc đến đời sống và sinh kế của người dân. Các hộ gia đình phải điều chỉnh thói quen sinh hoạt và phương thức sản xuất để thích ứng với tình hình mới.

### 💡 Góc Nhìn Giải Pháp

Từ góc độ giải pháp, {title} đòi hỏi cách tiếp cận tổng thể và đồng bộ từ nhiều phía. Các biện pháp ngắn hạn cần được triển khai ngay để ổn định tình hình, trong đó ưu tiên bảo vệ quyền lợi người dân.

Sự kiện {title} một lần nữa cho thấy tầm quan trọng của việc xây dựng hệ thống quản lý hiệu quả và nâng cao ý thức cộng đồng. Bài học kinh nghiệm từ sự việc này cần được ghi nhận và áp dụng để ngăn chặn những tình huống tương tự trong tương lai."""

        # Đảm bảo đủ số từ
        current_words = len(article.split())
        if current_words < min_words:
            additional_content = self.generate_additional_journalistic_content(title, category, min_words - current_words)
            article += "\n\n" + additional_content

        return article

    def generate_comprehensive_article(self, title, category, min_words):
        """Tạo bài báo toàn diện khi không có đủ dữ liệu"""
        return self.generate_fallback_journalistic_article(title, category, min_words)

    def generate_fallback_content(self, title, category, target_words):
        """Tạo nội dung fallback"""
        content = self.generate_fallback_journalistic_article(title, category, target_words)
        return content

    def generate_script_intro(self, title, category):
        """Tạo phần mở đầu script chi tiết"""
        intros = {
            'thời sự': [
                f"Kính thưa quý vị và các bạn! Trong chương trình ngày hôm nay, chúng ta sẽ cùng tìm hiểu sâu về vấn đề {title} - một chủ đề đang thu hút sự quan tâm đặc biệt của dư luận. Sự kiện này không chỉ tác động trực tiếp đến đời sống người dân mà còn đặt ra nhiều vấn đề quan trọng về công tác quản lý và phát triển. Với tư cách là những người làm truyền thông, chúng tôi nhận thấy cần phải cung cấp một cái nhìn toàn diện và khách quan về vấn đề này, từ những phân tích chuyên sâu đến những góc nhìn đa chiều từ các bên liên quan.",
                f"Thưa quý vị khán giả! {title} đang là tâm điểm chú ý trong những ngày qua, với những diễn biến phức tạp và nhiều hệ lụy đáng quan tâm. Trong bản tin đặc biệt hôm nay, chúng tôi sẽ mang đến cho quý vị một bức tranh tổng thể về sự việc, từ nguyên nhân, diễn biến cho đến những giải pháp đang được đề xuất. Đây không chỉ là một bản tin thông thường mà là một cuộc trao đổi chuyên sâu, nơi chúng ta cùng nhau tìm hiểu và phân tích mọi khía cạnh của vấn đề."
            ],
            'kinh doanh': [
                f"Kính chào quý vị và các bạn! Trong bối cảnh thị trường đang có nhiều biến động, {title} đã trở thành chủ đề nóng hổi được giới chuyên môn và cộng đồng doanh nghiệp đặc biệt quan tâm. Sự kiện này không chỉ ảnh hưởng đến hoạt động sản xuất kinh doanh mà còn tác động sâu rộng đến cấu trúc thị trường và chiến lược phát triển của các doanh nghiệp. Trong chương trình hôm nay, chúng ta sẽ cùng phân tích kỹ lưỡng mọi khía cạnh của vấn đề, từ tác động ngắn hạn đến xu hướng dài hạn, và đặc biệt là những bài học kinh nghiệm quý giá cho cộng đồng doanh nghiệp.",
                f"Thưa quý vị! Trong thế giới kinh doanh luôn biến động, {title} đã tạo ra những thay đổi đáng kể trong cách vận hành của thị trường. Sự kiện này không chỉ là câu chuyện của riêng một doanh nghiệp hay một ngành hàng, mà đã trở thành điểm nóng thu hút sự chú ý của toàn bộ nền kinhế. Với sự tham gia của các chuyên gia hàng đầu và những phân tích chi tiết, chúng tôi hy vọng sẽ mang đến cho quý vị một cái nhìn toàn diện và sâu sắc về vấn đề này."
            ],
            'giáo dục': [
                f"Kính thưa quý vị và các bạn! Trong chương trình giáo dục hôm nay, chúng ta sẽ cùng tìm hiểu sâu về vấn đề {title} - một chủ đề đang nhận được sự quan tâm đặc biệt từ phụ huynh, học sinh và các nhà giáo dục. Sự kiện này không chỉ tác động đến chất lượng giáo dục mà còn ảnh hưởng đến tương lai của thế hệ trẻ. Với tư cách là những người làm công tác giáo dục, chúng tôi nhận thấy cần phải cung cảm báo một cái nhìn toàn diện về vấn đề này, từ những phân tích chuyên sâu đến những giải pháp thiết thực.",
                f"Thưa quý vị khán giả! {title} đang là tâm điểm chú ý trong lĩnh vực giáo dục những ngày qua, với những diễn biến quan trọng và nhiều ý kiến đa chiều. Trong chương trình đặc biệt hôm nay, chúng tôi sẽ mang đến cho quý vị một cái nhìn tổng thể về sự việc, từ nguyên nhân, diễn biến cho đến những giải pháp đang được đề xuất. Đây là cơ hội để chúng ta cùng nhau tìm hiểu và phân tích mọi khía cạnh của vấn đề giáo dục này."
            ]
        }

        templates = intros.get(category, intros['thời sự'])
        return random.choice(templates)

    def generate_script_analysis(self, title, category):
        """Tạo phần phân tích cho script"""
        analysis_templates = {
            'thời sự': [
                f"Phân tích sâu về {title} cho thấy đây không chỉ là vấn đề đơn lẻ mà là biểu hiện của những thách thức lớn hơn trong quản lý đô thị và phát triển bền vững. Các chuyên gia nhấn mạnh rằng cần có cách tiếp cận toàn diện, kết hợp giữa giải pháp kỹ thuật và quản lý, đồng thời tăng cường sự tham gia của cộng đồng trong quá trình ra quyết định.",
                f"Từ góc độ chuyên môn, {title} đặt ra nhiều câu hỏi quan trọng về hiệu quả của các cơ chế quản lý hiện hành. Các chuyên gia cho rằng cần rà soát lại toàn bộ quy trình, từ khâu lập kế hoạch đến triển khai và giám sát, để đảm bảo tính đồng bộ và hiệu quả trong việc giải quyết vấn đề."
            ],
            'kinh doanh': [
                f"Phân tích chi tiết về {title} cho thấy tác động đa chiều đến cấu trúc thị trường và hành vi của các bên tham gia. Sự kiện này không chỉ ảnh hưởng đến hoạt động sản xuất kinh doanh trước mắt mà còn định hình lại xu hướng phát triển dài hạn của toàn ngành.",
                f"Từ góc độ chiến lược, {title} buộc các doanh nghiệp phải xem xét lại mô hình kinh doanh và khả năng thích ứng với biến động. Các chuyên gia khuyến nghị cần kết hợp giữa bảo tồn giá trị cốt lõi và đổi mới sáng tạo để duy trì lợi thế cạnh tranh trong bối cảnh mới."
            ],
            'giáo dục': [
                f"Phân tích chuyên sâu về {title} cho thấy những thách thức và cơ hội trong việc nâng cao chất lượng giáo dục. Sự kiện này nhấn mạnh sự cần thiết của việc đổi mới toàn diện, từ chương trình đào tạo đến phương pháp giảng dạy và cơ sở vật chất.",
                f"Từ góc độ phát triển bền vững, {title} đặt ra yêu cầu về việc xây dựng hệ thống giáo dục linh hoạt, có khả năng thích ứng với những thay đổi của xã hội và nhu cầu của thị trường lao động trong tương lai."
            ]
        }
        templates = analysis_templates.get(category, analysis_templates['thời sự'])
        return random.choice(templates)

    def generate_script_facts(self, title, category):
        """Tạo phần thông tin thực tế cho script"""
        facts_templates = {
            'thời sự': [
                f"Theo số liệu thống kê mới nhất, {title} đã tác động trực tiếp đến hàng nghìn hộ dân tại nhiều khu vực khác nhau. Các con số cho thấy mức độ nghiêm trọng của vấn đề và sự cần thiết phải có những biện pháp can thiệp kịp thời và hiệu quả.",
                f"Ghi nhận thực tế cho thấy {title} đã gây ra những thiệt hại đáng kể về kinh tế và xã hội. Các chuyên gia ước tính thiệt hại có thể lên đến hàng trăm tỷ đồng, chưa kể những tác động lâu dài đến đời sống người dân và môi trường."
            ],
            'kinh doanh': [
                f"Số liệu từ các cơ quan chức năng cho thấy {title} đã ảnh hưởng đến hoạt động của hàng trăm doanh nghiệp, với tổng thiệt hại ước tính lên đến hàng nghìn tỷ đồng. Các chỉ số kinh tế vĩ mô cũng phản ánh mức độ tác động sâu rộng của sự kiện này.",
                f"Theo báo cáo mới nhất, {title} đã dẫn đến sự sụt giảm đáng kể trong chỉ số niềm tin của nhà đầu tư và người tiêu dùng. Các chuyên gia cảnh báo về khả năng ảnh hưởng lâu dài đến tốc độ phục hồi và tăng trưởng của nền kinh tế."
            ],
            'giáo dục': [
                f"Theo thống kê từ Bộ Giáo dục và Đào tạo, {title} đã ảnh hưởng đến hàng triệu học sinh, sinh viên trên cả nước. Các con số cho thấy sự cần thiết của việc đầu tư và cải tiến toàn diện hệ thống giáo dục.",
                f"Khảo sát thực tế cho thấy {title} đã làm bộc lộ nhiều điểm yếu trong hệ thống giáo dục hiện hành. Các chuyên gia nhấn mạnh nhu cầu cấp thiết về việc đổi mới phương pháp giảng dạy và nâng cao chất lượng đào tạo."
            ]
        }
        templates = facts_templates.get(category, facts_templates['thời sự'])
        return random.choice(templates)

    def generate_script_reactions(self, title, category):
        """Tạo phần phản ứng và đánh giá cho script"""
        reactions_templates = {
            'thời sự': [
                f"Phản ứng từ cộng đồng trước {title} cho thấy sự quan tâm và lo ngại sâu sắc. Người dân tại các khu vực chịu ảnh hưởng đã lên tiếng kêu gọi sự vào cuộc kịp thời và quyết liệt của các cơ quan chức năng.",
                f"Các chuyên gia và tổ chức xã hội đã có nhiều ý kiến đánh giá về {title}. Đa số đều thống nhất về tính phức tạp của vấn đề và sự cần thiết phải có giải pháp đồng bộ, toàn diện."
            ],
            'kinh doanh': [
                f"Phản ứng từ cộng đồng doanh nghiệp trước {title} rất đa dạng, từ thận trọng đến chủ động thích ứng. Các hiệp hội ngành nghề đã tích cực phối hợp để tìm kiếm giải pháp chung và hỗ trợ các thành viên.",
                f"Các chuyên gia kinh tế đánh giá cao khả năng thích ứng của doanh nghiệp Việt Nam trước {title}. Tuy nhiên, họ cũng cảnh báo về những thách thức dài hạn cần được quan tâm giải quyết."
            ],
            'giáo dục': [
                f"Phản ứng từ phụ huynh và học sinh trước {title} cho thấy sự quan tâm đặc biệt đến chất lượng giáo dục. Nhiều ý kiến đề xuất cần có sự đổi mới toàn diện trong phương pháp giảng dạy và đánh giá.",
                f"Các chuyên gia giáo dục đánh giá {title} là cơ hội để nhìn nhận lại toàn bộ hệ thống. Họ nhấn mạnh sự cần thiết của việc kết hợp giữa kinh nghiệm quốc tế và đặc thù Việt Nam trong cải cách giáo dục."
            ]
        }
        templates = reactions_templates.get(category, reactions_templates['thời sự'])
        return random.choice(templates)

    def generate_script_solutions(self, title, category):
        """Tạo phần giải pháp và khuyến nghị cho script"""
        solutions_templates = {
            'thời sự': [
                f"Trước tình hình {title}, các chuyên gia đề xuất một loạt giải pháp đồng bộ. Về ngắn hạn, cần tập trung vào các biện pháp khẩn cấp để ổn định tình hình. Về dài hạn, cần xây dựng kế hoạch chiến lược với lộ trình rõ ràng, tập trung vào cải cách thể chế và nâng cao năng lực quản lý.",
                f"Giải pháp cho {title} cần kết hợp giữa yếu tố kỹ thuật và quản lý. Các chuyên gia nhấn mạnh tầm quan trọng của việc ứng dụng công nghệ tiên tiến, đồng thời hoàn thiện khung pháp lý và cơ chế phối hợp giữa các bên liên quan."
            ],
            'kinh doanh': [
                f"Để ứng phó với {title}, các chuyên gia kinh tế đề xuất nhiều giải pháp đa dạng. Về phía nhà nước, cần hoàn thiện khung pháp lý và chính sách hỗ trợ. Về phía doanh nghiệp, cần chủ động đổi mới mô hình kinh doanh và nâng cao năng lực cạnh tranh.",
                f"Giải pháp cho {title} cần tiếp cận từ cả góc độ vĩ mô và vi mô. Các chuyên gia khuyến nghị cần kết hợp giữa biện pháp tình thế và chiến lược dài hạn, đồng thời tăng cường sự phối hợp giữa các bên trong việc thực hiện các giải pháp."
            ],
            'giáo dục': [
                f"Trước thách thức từ {title}, các chuyên gia giáo dục đề xuất nhiều giải pháp toàn diện. Về chương trình đào tạo, cần cập nhật và điều chỉnh phù hợp với xu hướng phát triển. Về phương pháp giảng dạy, cần đổi mới theo hướng phát huy tính chủ động và sáng tạo của người học.",
                f"Giải pháp cho {title} cần tập trung vào việc xây dựng hệ thống giáo dục linh hoạt và bền vững. Các chuyên gia nhấn mạnh tầm quan trọng của việc đầu tư vào cơ sở vật chất, nâng cao năng lực đội ngũ giảng viên và ứng dụng công nghệ trong giảng dạy."
            ]
        }
        templates = solutions_templates.get(category, solutions_templates['thời sự'])
        return random.choice(templates)

    def generate_script_conclusion(self, title, category):
        """Tạo phần kết luận cho script"""
        conclusion_templates = {
            'thời sự': [
                f"Sự kiện {title} một lần nữa cho thấy tầm quan trọng của việc xây dựng hệ thống quản lý hiệu quả và nâng cao ý thức cộng đồng. Bài học kinh nghiệm từ sự việc này cần được ghi nhận và áp dụng để ngăn chặn những tình huống tương tự trong tương lai. Người dân và chính quyền cần tiếp tục phối hợp chặt chẽ để tìm ra giải pháp bền vững.",
                f"Trước mắt, các biện pháp xử lý {title} đang được triển khai quyết liệt. Tuy nhiên, về lâu dài, cần có sự thay đổi căn bản trong cách tiếp cận và giải quyết vấn đề. Sự chung tay của toàn xã hội sẽ tạo ra sức mạnh tổng hợp để vượt qua thách thức và xây dựng cộng đồng ngày càng tốt đẹp hơn."
            ],
            'kinh doanh': [
                f"{title} đang định hình lại bức tranh kinh doanh và đầu tư. Các doanh nghiệp cần nhanh chóng thích ứng với xu hướng mới và tìm kiếm cơ hội trong thách thức. Sự linh hoạt và sáng tạo sẽ là yếu tố then chốt để thành công trong giai đoạn chuyển đổi quan trọng này.",
                f"Bối cảnh kinh doanh sau {title} sẽ có nhiều thay đổi so với trước đây. Các doanh nghiệp cần chuẩn bị cho những kịch bản phát triển mới và xây dựng năng lực cạnh tranh bền vững. Việc nắm bắt cơ hội và quản lý rủi ro hiệu quả sẽ quyết định vị thế của doanh nghiệp trong tương lai."
            ],
            'giáo dục': [
                f"{title} đang định hình lại tương lai của giáo dục và đào tạo. Các cơ sở giáo dục cần nhanh chóng thích ứng với xu hướng mới và tìm kiếm cơ hội đổi mới. Sự sáng tạo và đổi mới sẽ là yếu tố then chốt để nâng cao chất lượng giáo dục trong giai đoạn phát triển quan trọng này.",
                f"Bối cảnh giáo dục sau {title} sẽ có nhiều thay đổi so với trước đây. Các cơ sở giáo dục cần chuẩn bị cho những mô hình đào tạo mới và xây dựng năng lực cạnh tranh bền vững. Việc nắm bắt cơ hội và ứng phó với thách thức hiệu quả sẽ quyết định chất lượng giáo dục trong tương lai."
            ]
        }
        templates = conclusion_templates.get(category, conclusion_templates['thời sự'])
        return random.choice(templates)

    def generate_additional_script_content(self, title, category, word_needed):
        """Tạo nội dung script bổ sung"""
        additional_content = ""
        current_words = 0

        while current_words < word_needed:
            paragraph = self.generate_additional_perspective(title, category)
            additional_content += paragraph + "\n\n"
            current_words += len(paragraph.split())

        return additional_content.strip()

    def generate_fallback_script_content(self, title, category):
        """Tạo nội dung script fallback - PHƯƠNG THỨC BỊ THIẾU ĐÃ ĐƯỢC THÊM"""
        try:
            current_date = self.get_current_time().split()[0]

            script_content = f"""Kính thưa quý vị và các bạn!

Hôm nay, {current_date}, chúng ta sẽ cùng tìm hiểu sâu về vấn đề {title} - một chủ đề đang thu hút sự quan tâm đặc biệt của dư luận.

Theo các chuyên gia, {title} không chỉ là vấn đề trước mắt mà còn đặt ra nhiều thách thức lâu dài về quản lý và phát triển. Sự kiện này đã tác động trực tiếp đến đời sống người dân và hoạt động của nhiều lĩnh vực khác nhau.

Phân tích chuyên sâu cho thấy cần có cách tiếp cận toàn diện để giải quyết vấn đề này, kết hợp giữa giải pháp kỹ thuật và quản lý, đồng thời tăng cường sự tham gia của cộng đồng.

Trên thực tế, {title} đã gây ra những thiệt hại đáng kể về kinh tế và xã hội. Các số liệu thống kê cho thấy mức độ nghiêm trọng của vấn đề và sự cần thiết phải có những biện pháp can thiệp kịp thời.

Phản ứng từ cộng đồng cho thấy sự quan tâm và lo ngại sâu sắc. Người dân và các bên liên quan đã lên tiếng kêu gọi sự vào cuộc quyết liệt của các cơ quan chức năng.

Trước tình hình này, các chuyên gia đề xuất một loạt giải pháp đồng bộ. Về ngắn hạn, cần tập trung vào các biện pháp khẩn cấp để ổn định tình hình. Về dài hạn, cần xây dựng kế hoạch chiến lược với lộ trình rõ ràng.

{title} một lần nữa cho thấy tầm quan trọng của việc xây dựng hệ thống quản lý hiệu quả và nâng cao ý thức cộng đồng. Bài học kinh nghiệm từ sự việc này cần được ghi nhận và áp dụng để ngăn chặn những tình huống tương tự trong tương lai.

Xin cảm ơn quý vị và các bạn đã theo dõi!"""

            return script_content

        except Exception as e:
            print(f"❌ Lỗi tạo script fallback: {e}")
            return f"""Kính thưa quý vị và các bạn!

Hôm nay chúng ta sẽ cùng thảo luận về {title}. Đây là một chủ đề quan trọng đang được dư luận đặc biệt quan tâm.

Vấn đề này có nhiều khía cạnh phức tạp cần được phân tích kỹ lưỡng. Chúng tôi sẽ mang đến cho quý vị những thông tin mới nhất và các góc nhìn chuyên sâu.

Xin cảm ơn quý vị đã theo dõi chương trình!"""

    def clean_text(self, text):
        """Làm sạch văn bản"""
        if not text:
            return ""
        # Loại bỏ HTML tags
        text = re.sub(r'<[^>]+>', '', text)
        # Loại bỏ URL
        text = re.sub(r'http\S+', '', text)
        # Loại bỏ khoảng trắng thừa
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    def get_current_time(self):
        """Lấy thời gian hiện tại"""
        vietnam_tz = timezone(timedelta(hours=7))
        return datetime.now(vietnam_tz).strftime('%d/%m/%Y %H:%M:%S')

    # ==============================================
    # PHƯƠNG THỨC QUAN TRỌNG ĐÃ SỬA - LOẠI BỎ LẶP ĐOẠN VÀ LẶP Ý
    # ==============================================

    def generate_detailed_script_content(self, title, category):
        """Tạo nội dung script chi tiết > 700 từ - ĐÃ SỬA LỖI LẶP ĐOẠN VÀ LẶP Ý"""
        try:
            # Tạo danh sách các phần nội dung độc đáo
            content_sections = []
            
            # Phần mở đầu - chỉ gọi 1 lần
            intro = self.generate_script_intro(title, category)
            content_sections.append(intro)
            
            # Phân tích chuyên sâu - chỉ gọi 1 lần  
            analysis = self.generate_script_analysis(title, category)
            content_sections.append(analysis)
            
            # Thông tin thực tế - chỉ gọi 1 lần
            facts = self.generate_script_facts(title, category)
            content_sections.append(facts)
            
            # Phản ứng và đánh giá - chỉ gọi 1 lần
            reactions = self.generate_script_reactions(title, category)
            content_sections.append(reactions)
            
            # Giải pháp và khuyến nghị - chỉ gọi 1 lần
            solutions = self.generate_script_solutions(title, category)
            content_sections.append(solutions)
            
            # Kết luận - chỉ gọi 1 lần
            conclusion = self.generate_script_conclusion(title, category)
            content_sections.append(conclusion)

            # Kết hợp tất cả các phần
            full_content = "\n\n".join(content_sections)

            # Đảm bảo đủ 700 từ bằng cách thêm nội dung độc đáo nếu cần
            current_words = len(full_content.split())
            if current_words < 700:
                # Tạo nội dung bổ sung độc đáo từ các góc nhìn khác
                additional_needed = 700 - current_words
                additional_content = self._generate_unique_additional_content(title, category, additional_needed)
                full_content += "\n\n" + additional_content

            return full_content

        except Exception as e:
            print(f"❌ Lỗi tạo nội dung script: {e}")
            return self.generate_fallback_script_content(title, category)

    def _generate_unique_additional_content(self, title, category, word_needed):
        """Tạo nội dung bổ sung độc đáo không trùng lặp"""
        additional_perspectives = [
            f"Xét từ góc độ lịch sử, {title} cho thấy những bài học quan trọng về quá trình phát triển và những thách thức cần vượt qua. Kinh nghiệm từ các giai đoạn trước đây có thể cung cấp những gợi ý quý giá cho việc giải quyết vấn đề hiện tại.",
            
            f"Về mặt công nghệ và đổi mới sáng tạo, {title} mở ra cơ hội ứng dụng các giải pháp tiên tiến. Sự phát triển của công nghệ số và trí tuệ nhân tạo có thể đóng góp tích cực vào việc tìm kiếm các phương án tối ưu.",
            
            f"Trên bình diện quốc tế, {title} cũng được quan tâm và theo dõi sát sao. Các nước trong khu vực và trên thế giới có thể có những cách tiếp cận khác nhau, từ đó rút ra những kinh nghiệm hữu ích cho bối cảnh Việt Nam.",
            
            f"Về tác động đến các nhóm yếu thế trong xã hội, {title} cần được xem xét một cách toàn diện. Các chính sách hỗ trợ cần được thiết kế phù hợp để đảm bảo không ai bị bỏ lại phía sau.",
            
            f"Xét về tính bền vững lâu dài, {title} đòi hỏi những giải pháp có tầm nhìn xa. Các quyết định hiện tại cần tính đến tác động trong tương lai 5-10 năm tới."
        ]
        
        # Chọn ngẫu nhiên các góc nhìn cho đến khi đủ số từ
        selected_perspectives = []
        current_words = 0
        available_perspectives = additional_perspectives.copy()
        
        while current_words < word_needed and available_perspectives:
            # Chọn ngẫu nhiên một góc nhìn
            perspective = random.choice(available_perspectives)
            selected_perspectives.append(perspective)
            current_words += len(perspective.split())
            
            # Loại bỏ góc nhìn đã chọn để tránh lặp lại
            available_perspectives.remove(perspective)
            
            # Nếu đã hết góc nhìn, thoát khỏi vòng lặp
            if not available_perspectives:
                break
        
        return "\n\n".join(selected_perspectives)

# ==============================================
# LỚP CHATBOT CHÍNH - ĐÃ SỬA ĐỂ KẾT NỐI VỚI ADMIN
# ==============================================

class SmartVideoNewsChatbot:
    def __init__(self):
        self.data_manager = DataManager()
        self.media_processor = AdvancedMediaProcessor()
        self.content_generator = ContentGenerator()

        # Trạng thái hiện tại
        self.current_video = None
        self.current_category = None
        self.current_article = ""
        self.current_script = ""
        self.digest_videos = []  # Thêm danh sách video tổng hợp

        # Cache và lịch sử
        self.rss_cache = {}
        self.last_rss_update = {}
        self.user_history = []

        # Tải dữ liệu TỪ ADMIN - ĐÃ SỬA
        self.refresh_data_from_admin()

        # Quản lý lịch sử - MỚI THÊM
        self.history_manager = ChatHistoryManager()

        print("✅ Chatbot đã được khởi tạo thành công với dữ liệu từ Admin")

    def refresh_data_from_admin(self):
        """Làm mới dữ liệu từ admin - PHƯƠNG THỨC MỚI QUAN TRỌNG"""
        try:
            print("🔄 Đang đồng bộ dữ liệu từ Admin...")
            self.rss_feeds, self.youtube_channels = self.data_manager.refresh_data()
            print(f"✅ Đã đồng bộ: {len(self.rss_feeds)} RSS feeds, {len(self.youtube_channels)} YouTube channels")
        except Exception as e:
            print(f"❌ Lỗi đồng bộ dữ liệu từ Admin: {e}")

    def get_vietnam_time(self):
        """Lấy thời gian Việt Nam"""
        vietnam_tz = timezone(timedelta(hours=7))
        return datetime.now(vietnam_tz)

    def format_time(self, dt=None):
        """Định dạng thời gian"""
        if dt is None:
            dt = self.get_vietnam_time()
        return dt.strftime('%d/%m/%Y %H:%M:%S')

    def detect_category(self, query):
        """Phát hiện category từ query - ĐÃ CẢI THIỆN VỚI TỪ KHÓA MỞ RỘNG"""
        query_lower = query.lower()

        # Sử dụng hệ thống từ khóa mở rộng
        for category, keywords in EXPANDED_KEYWORDS.items():
            if any(keyword in query_lower for keyword in keywords):
                print(f"🎯 Tìm thấy category '{category}' với từ khóa: {query}")
                return category

        # Kiểm tra tên category trực tiếp
        for category in self.get_all_categories():
            if category in query_lower:
                return category

        print(f"🔍 Không tìm thấy category phù hợp cho: {query}, sử dụng mặc định 'thời sự'")
        return 'thời sự'  # Mặc định

    def evaluate_deep_relevance(self, user_query, video_title, video_description, video_content, video_category=None):
        """Đánh giá mức độ liên quan sâu giữa yêu cầu và nội dung video - ĐÃ SỬA THEO YÊU CẦU"""
        try:
            if not user_query:
                return "🔍 Không có thông tin để đánh giá"

            query_lower = user_query.lower()
            title_lower = video_title.lower()
            desc_lower = video_description.lower()
            content_lower = video_content.lower() if video_content else ""

            # VẤN ĐỀ 1 ĐÃ SỬA: ĐIỀU CHỈNH ĐÁNH GIÁ % PHÙ HỢP
            # TRƯỜNG HỢP ĐẶC BIỆT: Nếu user_query trùng với video_category (chủ đề của video)
            if video_category and query_lower == video_category.lower():
                return "🎯 Rất phù hợp (98-100%) - Video được chọn trực tiếp từ chủ đề này"

            # TRƯỜNG HỢP ĐẶC BIỆT: Nếu video được tạo trực tiếp từ chủ đề (có video_category)
            if video_category and video_category.lower() in query_lower:
                return "✅ Phù hợp (80-95%) - Video được tạo trực tiếp từ chủ đề"

            # TRƯỜNG HỢP ĐẶC BIỆT: Nếu không có query cụ thể (tạo bài từ chủ đề)
            if query_lower in ['', 'none', 'null']:
                return "🎯 Rất phù hợp (95-98%) - Video phù hợp với chủ đề đã chọn"

            # PHÂN TÍCH CHI TIẾT TỪ KHÓA - ĐÃ CẢI THIỆN
            query_words = set()
            for word in query_lower.split():
                if len(word) > 2:  # Bỏ qua các từ quá ngắn
                    query_words.add(word)

            # Thêm từ khóa mở rộng dựa trên ngữ cảnh
            expanded_keywords = self.expand_keywords(query_words)
            query_words.update(expanded_keywords)

            # Tính điểm cho tiêu đề - TĂNG TRỌNG SỐ CHO SỰ PHÙ HỢP CHÍNH XÁC
            title_score = 0
            title_words = title_lower.split()
            exact_match_found = False

            for keyword in query_words:
                # Kiểm tra khớp chính xác từ khóa quan trọng
                if (keyword in title_lower and
                    len(keyword) > 3 and  # Từ khóa đủ dài
                    any(edu_word in keyword for edu_word in ['tuyển', 'sinh', 'đại', 'học', 'thi', 'trung', 'phổ', 'thông']) if 'tuyển sinh' in query_lower else True):

                    # Ưu tiên từ khóa xuất hiện sớm trong tiêu đề
                    if title_words and keyword in title_words[0]:  # Từ đầu tiên
                        title_score += 8
                        exact_match_found = True
                    elif len(title_words) > 1 and keyword in title_words[1]:  # Từ thứ hai
                        title_score += 6
                        exact_match_found = True
                    else:
                        title_score += 4
                elif keyword in title_lower:
                    title_score += 2

            # Tính điểm cho mô tả
            desc_score = 0
            for keyword in query_words:
                if keyword in desc_lower:
                    # Ưu tiên từ khóa quan trọng trong mô tả
                    if len(keyword) > 4:
                        desc_score += 3
                    else:
                        desc_score += 1

            # Tính điểm cho nội dung
            content_score = 0
            for keyword in query_words:
                if keyword in content_lower:
                    content_score += 1

            # ĐIỀU CHỈNH ĐIỂM SỐ DỰA TRÊN ĐỘ PHÙ HỢP THỰC TẾ
            total_score = title_score + desc_score + content_score

            # Điểm tối đa có thể đạt được (điều chỉnh theo độ khó)
            base_max_score = len(query_words) * 12  # 8+3+1 = 12

            if base_max_score == 0:
                return "❌ Không đủ thông tin để đánh giá"

            # ĐIỀU CHỈNH QUAN TRỌNG: Giảm điểm số chung để phản ánh đúng thực tế
            relevance_percentage = (total_score / base_max_score) * 100

            # ĐIỀU CHỈNH NGƯỠNG ĐÁNH GIÁ THEO YÊU CẦU
            if exact_match_found and relevance_percentage >= 70:
                return f"🎯 Rất phù hợp ({min(98, relevance_percentage):.1f}%) - Khớp chính xác từ khóa quan trọng"
            elif relevance_percentage >= 80:
                return f"🎯 Rất phù hợp ({min(98, relevance_percentage):.1f}%)"
            elif relevance_percentage >= 60:
                return f"✅ Phù hợp ({relevance_percentage:.1f}%)"
            elif relevance_percentage >= 40:
                return f"⚠️ Khá phù hợp ({relevance_percentage:.1f}%)"
            elif relevance_percentage >= 20:
                return f"📌 Tương đối phù hợp ({relevance_percentage:.1f}%)"
            elif relevance_percentage >= 5:
                return f"🔍 Ít phù hợp ({relevance_percentage:.1f}%)"
            else:
                return f"❌ Không phù hợp ({relevance_percentage:.1f}%)"

        except Exception as e:
            return f"🔍 Lỗi đánh giá: {str(e)}"

    def expand_keywords(self, keywords):
        """Mở rộng từ khóa dựa trên ngữ cảnh - ĐÃ CẢI THIỆN"""
        expanded = set()

        # Sử dụng hệ thống từ khóa mở rộng
        for category, category_keywords in EXPANDED_KEYWORDS.items():
            for keyword in keywords:
                if keyword in category_keywords:
                    # Thêm tất cả từ khóa liên quan trong category đó
                    expanded.update(category_keywords)
                    break

        # Thêm các biến thể của từ khóa
        for keyword in keywords:
            expanded.add(keyword)
            # Thêm các từ đồng nghĩa/thuộc cùng lĩnh vực
            if 'ai' in keyword or 'trí tuệ nhân tạo' in keyword:
                expanded.update(['machine learning', 'deep learning', 'neural network', 'artificial intelligence'])
            elif 'robot' in keyword:
                expanded.update(['robotics', 'automation', 'tự động hóa', 'công nghệ robot'])
            elif 'công nghệ' in keyword:
                expanded.update(['tech', 'technology', 'kỹ thuật', 'đổi mới sáng tạo'])

        return expanded

    def get_all_categories(self):
        """Lấy tất cả categories - ĐÃ SỬA ĐỂ BAO GỒM TẤT CẢ CHỦ ĐỀ TỪ ADMIN"""
        # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI LẤY
        self.refresh_data_from_admin()

        rss_categories = list(self.rss_feeds.keys())
        youtube_categories = list(self.youtube_channels.keys())
        all_categories = rss_categories + youtube_categories  # ĐÃ SỬA LỖI Ở ĐÂY

        # THÊM CÁC CHỦ ĐỀ MỚI TỪ EXPANDED_KEYWORDS ĐỂ ĐẢM BẢO TÌM KIẾM TẤT CẢ
        for category in EXPANDED_KEYWORDS.keys():
            if category not in all_categories:
                all_categories.append(category)

        return all_categories

    def validate_rss_url(self, url):
        """Validate RSS URL"""
        try:
            feed = feedparser.parse(url)
            return len(feed.entries) > 0
        except:
            return False

    def validate_youtube_channel(self, channel_id):
        """Validate YouTube channel"""
        try:
            rss_url = f"https://www.youtube.com/feeds/videos.xml?channel_id={channel_id}"
            feed = feedparser.parse(rss_url)
            return len(feed.entries) > 0
        except:
            return False

    def get_youtube_videos(self, channel_id, count=10):
        """Lấy video từ YouTube channel"""
        try:
            rss_url = f"https://www.youtube.com/feeds/videos.xml?channel_id={channel_id}"
            feed = feedparser.parse(rss_url)

            videos = []
            for entry in feed.entries[:count]:
                try:
                    video_id = entry.yt_videoid
                    video_data = {
                        'title': entry.title,
                        'description': getattr(entry, 'description', ''),
                        'link': f"https://www.youtube.com/watch?v={video_id}",
                        'thumbnail': f"https://img.youtube.com/vi/{video_id}/hqdefault.jpg",
                        'published': getattr(entry, 'published', self.format_time()),
                        'video_id': video_id,
                        'source': 'youtube'
                    }
                    videos.append(video_data)
                except Exception as e:
                    continue

            return videos
        except Exception as e:
            print(f"❌ Lỗi lấy video YouTube: {e}")
            return []

    def get_rss_videos(self, category, count=10):
        """Lấy video từ RSS feed - ĐÃ SỬA LỖI ẢNH"""
        try:
            # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI LẤY
            self.refresh_data_from_admin()

            url = self.rss_feeds.get(category)
            if not url:
                return []

            feed = feedparser.parse(url)
            videos = []

            for entry in feed.entries[:count*2]:
                try:
                    # Lấy thông tin cơ bản
                    title = getattr(entry, 'title', 'Không có tiêu đề')
                    description = getattr(entry, 'description', 'Không có mô tả')
                    link = getattr(entry, 'link', '#')
                    published = getattr(entry, 'published', self.format_time())

                    # TRÍCH XUẤT ẢNH TỪ RSS - SỬ DỤNG PHƯƠNG THỨC MỚI
                    thumbnail = self.media_processor.extract_image_from_rss(entry)

                    # Thử trích xuất frame từ video (chỉ với link hợp lệ)
                    extracted_frame = None
                    if link and ('youtube.com' in link or 'youtu.be' in link):
                        try:
                            extracted_frame = self.media_processor.extract_frame_from_video(link)
                        except Exception as e:
                            print(f"⚠️ Không thể trích xuất frame từ {link}: {e}")

                    video_data = {
                        'title': title,
                        'description': description,
                        'link': link,
                        'thumbnail': thumbnail,
                        'extracted_frame': extracted_frame,
                        'published': published,
                        'source': 'rss'
                    }
                    videos.append(video_data)
                except Exception as e:
                    print(f"❌ Lỗi xử lý entry RSS: {e}")
                    continue

            return videos
        except Exception as e:
            print(f"❌ Lỗi lấy video RSS: {e}")
            return []

    def get_videos(self, category, count=10):
        """Lấy video từ cả RSS và YouTube - ĐÃ SỬA ĐỂ TÌM TẤT CẢ CHỦ ĐỀ TỪ ADMIN"""
        # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI TÌM KIẾM
        self.refresh_data_from_admin()

        videos = []

        # Lấy từ RSS - SỬA: TÌM TẤT CẢ CHỦ ĐỀ CÓ LIÊN QUAN
        for cat in self.rss_feeds.keys():
            if category.lower() in cat.lower() or cat.lower() in category.lower():
                print(f"🔍 Tìm video từ RSS: {cat}")
                videos.extend(self.get_rss_videos(cat, count))

        # Lấy từ YouTube - SỬA: TÌM TẤT CẢ CHỦ ĐỀ CÓ LIÊN QUAN
        for cat in self.youtube_channels.keys():
            if category.lower() in cat.lower() or cat.lower() in category.lower():
                print(f"🔍 Tìm video từ YouTube: {cat}")
                channel_id = self.youtube_channels[cat]
                videos.extend(self.get_youtube_videos(channel_id, count))

        # Nếu không tìm thấy video, thử tìm trong tất cả các chủ đề có từ khóa liên quan
        if not videos:
            print(f"🔍 Không tìm thấy video trực tiếp, đang tìm kiếm mở rộng...")
            # Tìm các chủ đề có từ khóa liên quan trong EXPANDED_KEYWORDS
            for cat, keywords in EXPANDED_KEYWORDS.items():
                if any(keyword in category.lower() for keyword in keywords):
                    print(f"🔍 Tìm video mở rộng từ: {cat}")

                    # Thử RSS
                    if cat in self.rss_feeds:
                        videos.extend(self.get_rss_videos(cat, 3))

                    # Thử YouTube
                    if cat in self.youtube_channels:
                        channel_id = self.youtube_channels[cat]
                        videos.extend(self.get_youtube_videos(channel_id, 3))

        # Sắp xếp theo thời gian (mới nhất trước)
        def parse_published_date(pub_date):
            try:
                return datetime.strptime(pub_date, '%a, %d %b %Y %H:%M:%S %Z')
            except:
                try:
                    return datetime.strptime(pub_date, '%Y-%m-%dT%H:%M:%S%z')
                except:
                    return datetime.now()

        videos.sort(key=lambda x: parse_published_date(x['published']), reverse=True)

        return videos[:count]

    def get_youtube_transcript(self, video_id):
        """Lấy transcript từ YouTube"""
        try:
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['vi', 'en'])
            transcript_text = " ".join([entry['text'] for entry in transcript_list])
            return transcript_text
        except:
            return None

    def advanced_video_search(self, user_query, max_results=10):
        """Tìm kiếm video nâng cao dựa trên yêu cầu phức tạp - ĐÃ SỬA ĐỂ TÌM TẤT CẢ CHỦ ĐỀ"""
        try:
            print(f"🔍 Đang tìm kiếm nâng cao: {user_query}")

            # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI TÌM KIẾM
            self.refresh_data_from_admin()

            # Phân tích yêu cầu để tìm category phù hợp nhất - SỬA: TÌM TẤT CẢ CHỦ ĐỀ LIÊN QUAN
            target_categories = self.find_all_related_categories(user_query)
            print(f"🎯 Các categories được chọn: {target_categories}")

            # Tìm kiếm trong tất cả các category liên quan
            all_videos = []

            for category in target_categories:
                print(f"🔍 Đang tìm trong category: {category}")
                videos = self.get_videos(category, 8)
                for video in videos:
                    video['search_category'] = category
                    # THÊM ĐÁNH GIÁ PHÙ HỢP NGAY KHI TÌM KIẾM
                    base_content = video['description']
                    if video.get('source') == 'youtube' and video.get('video_id'):
                        transcript = self.get_youtube_transcript(video['video_id'])
                        if transcript:
                            base_content = transcript

                    relevance = self.evaluate_deep_relevance(
                        user_query,
                        video['title'],
                        video['description'],
                        base_content,
                        video['search_category']  # Thêm video_category
                    )
                    video['relevance_score'] = relevance
                    all_videos.append(video)

            # Sắp xếp theo độ phù hợp
            def get_relevance_score(relevance_str):
                try:
                    # Trích xuất phần trăm từ chuỗi "Rất phù hợp (95.5%)"
                    match = re.search(r'\(([\d.]+)%\)', relevance_str)
                    if match:
                        return float(match.group(1))
                    return 0
                except:
                    return 0

            all_videos.sort(key=lambda x: get_relevance_score(x.get('relevance_score', '0%')), reverse=True)

            if not all_videos:
                return self.handle_no_videos_found(user_query)

            # Lấy video phù hợp nhất
            best_video = all_videos[0]
            self.current_video = best_video
            self.current_category = best_video['search_category']

            print(f"✅ Đã tìm thấy video phù hợp: {best_video['title']}")
            print(f"📊 Độ phù hợp: {best_video.get('relevance_score', 'Không xác định')}")

            # Tạo nội dung bài báo
            return self.create_article_from_video(best_video, user_query)

        except Exception as e:
            print(f"❌ Lỗi tìm kiếm nâng cao: {str(e)}")
            return f"❌ Lỗi tìm kiếm: {str(e)}"

    def find_best_category_match(self, user_query):
        """Tìm category phù hợp nhất với yêu cầu - ĐÃ MỞ RỘNG VỚI NGUỒN MỚI"""
        query_lower = user_query.lower()

        # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI TÌM KIẾM
        self.refresh_data_from_admin()

        # Sử dụng hệ thống từ khóa mở rộng để tìm category phù hợp nhất
        best_category = 'thời sự vnexpress'  # Mặc định
        best_score = 0

        for category in self.get_all_categories():
            score = 0
            category_lower = category.lower()

            # Điểm cho category trực tiếp
            if category_lower in query_lower:
                score += 3

            # Điểm cho từ khóa liên quan trong hệ thống từ khóa mở rộng
            if category in EXPANDED_KEYWORDS:
                category_keywords = EXPANDED_KEYWORDS[category]
                for keyword in category_keywords:
                    if keyword in query_lower:
                        score += 2
                        # Ưu tiên từ khóa dài và cụ thể
                        if len(keyword) > 5:
                            score += 1

            if score > best_score:
                best_score = score
                best_category = category

        print(f"🎯 Category phù hợp nhất: {best_category} (điểm: {best_score})")
        return best_category

    def find_all_related_categories(self, user_query):
        """Tìm tất cả các categories liên quan đến yêu cầu - PHƯƠNG THỨC MỚI QUAN TRỌNG"""
        query_lower = user_query.lower()
        related_categories = set()

        # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI TÌM KIẾM
        self.refresh_data_from_admin()

        # Tìm tất cả categories có từ khóa liên quan
        for category, keywords in EXPANDED_KEYWORDS.items():
            for keyword in keywords:
                if keyword in query_lower:
                    related_categories.add(category)
                    break

        # Thêm categories có tên trực tiếp khớp
        for category in self.get_all_categories():
            if category.lower() in query_lower:
                related_categories.add(category)

        # Nếu không tìm thấy categories nào, sử dụng categories mặc định
        if not related_categories:
            related_categories = {'thời sự vnexpress', 'công nghệ', 'thời sự', 'giáo dục vnexpress'}

        print(f"🔍 Tìm thấy {len(related_categories)} categories liên quan: {related_categories}")
        return list(related_categories)

    def rank_videos_by_relevance(self, videos, user_query):
        """Xếp hạng video dựa trên độ phù hợp với yêu cầu"""
        if not videos:
            return []

        query_lower = user_query.lower()
        query_words = set(word for word in query_lower.split() if len(word) > 2)

        scored_videos = []
        for video in videos:
            score = 0

            # Kiểm tra tiêu đề
            title_lower = video['title'].lower()
            for word in query_words:
                if word in title_lower:
                    score += 3

            # Kiểm tra mô tả
            desc_lower = video.get('description', '').lower()
            for word in query_words:
                if word in desc_lower:
                    score += 2

            # Ưu tiên video mới
            try:
                published = video.get('published', '')
                if '2024' in published or '2025' in published:
                    score += 2
            except:
                pass

            # Ưu tiên video từ nguồn uy tín
            source = video.get('source', '')
            if source == 'youtube':
                score += 1

            if score > 0:
                video['relevance_score'] = score
                scored_videos.append(video)

        # Sắp xếp theo điểm giảm dần
        scored_videos.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)
        return scored_videos[:5]  # Trả về 5 video phù hợp nhất

    def handle_no_videos_found(self, user_query):
        """Xử lý khi không tìm thấy video phù hợp"""
        # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI THỬ CÁC CATEGORY KHÁC
        self.refresh_data_from_admin()

        # Thử các category khác nhau
        alternative_categories = ['thời sự vnexpress', 'VTV24', 'thời sự thanh niên', 'giáo dục vnexpress', 'công nghệ']

        for category in alternative_categories:
            videos = self.get_videos(category, 3)
            if videos:
                self.current_video = videos[0]
                self.current_category = category

                # Tạo bài báo với thông báo đặc biệt
                return self.create_fallback_article(user_query, category)

        # Nếu vẫn không tìm thấy
        return f"""❌ Không tìm thấy video phù hợp với yêu cầu: "{user_query}"

**🔍 Gợi ý:**
- Thử sử dụng từ khóa cụ thể hơn
- Kiểm tra lại chính tả
- Thử các chủ đề khác như: thời sự, thể thao, công nghệ
- Hoặc yêu cầu "tổng hợp video" để xem tất cả video có sẵn"""

    def create_fallback_article(self, user_query, category):
        """Tạo bài báo fallback khi không tìm thấy video hoàn toàn phù hợp"""
        videos = self.get_videos(category, 1)
        if not videos:
            return self.handle_no_videos_found(user_query)

        self.current_video = videos[0]
        self.current_category = category

        # Lấy transcript nếu là YouTube
        base_content = self.current_video['description']
        if self.current_video.get('source') == 'youtube' and self.current_video.get('video_id'):
            transcript = self.get_youtube_transcript(self.current_video['video_id'])
            if transcript:
                base_content = transcript

        # Tạo nội dung bài báo theo cấu trúc mới (600+ từ)
        article_content = self.content_generator.generate_article_content(
            self.current_video['title'],
            base_content,
            category,
            600
        )

        # Đánh giá mức độ phù hợp
        relevance_evaluation = self.evaluate_deep_relevance(
            user_query,
            self.current_video['title'],
            self.current_video['description'],
            base_content,
            self.current_category  # Thêm video_category
        )

        image_html = self.generate_image_html()

        # Tạo response đặc biệt
        icon = CATEGORY_ICONS.get(category, '📰')
        cleaned_title = self.content_generator.clean_title_for_content(self.current_video['title'])

        response = f"""## ⚠️ THÔNG TIN LIÊN QUAN

**🔍 Kết quả tìm kiếm cho:** "{user_query}"
{relevance_evaluation}

## {icon} {cleaned_title}

**📅 Ngày đăng:** {self.current_video['published']}
**🔍 Nguồn:** {self.current_video.get('source', 'unknown').upper()}

{image_html}

### 📝 Bài báo tham khảo:
{article_content}

---
**💡 Lưu ý:** Đây là video có nội dung gần nhất với yêu cầu của bạn.
**📺 Xem video gốc:** [{self.current_video['link']}]({self.current_video['link']})"""

        self.current_article = response
        return response

    def smart_content_creation(self, user_query, category):
        """Tạo nội dung thông minh dựa trên yêu cầu - ĐÃ CẢI THIỆN"""
        try:
            # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI TẠO NỘI DUNG
            self.refresh_data_from_admin()

            # Kiểm tra yêu cầu đặc biệt về giáo dục/tuyển sinh
            if any(keyword in user_query.lower() for keyword in ['tuyển sinh', 'đại học', 'cao đẳng', 'giáo dục']):
                print("🎓 Phát hiện yêu cầu giáo dục - tìm kiếm chuyên sâu...")
                return self.handle_education_request(user_query)

            # Thử tìm kiếm nâng cao trước
            advanced_result = self.advanced_video_search(user_query)
            if "Không tìm thấy" not in advanced_result and "Lỗi" not in advanced_result:
                return advanced_result

            # Nếu không thành công, dùng phương pháp thông thường
            return self.create_article_with_image(category, user_query)

        except Exception as e:
            # Fallback về phương pháp cơ bản
            return self.create_article_with_image(category, user_query)

    def handle_education_request(self, user_query):
        """Xử lý yêu cầu về giáo dục và tuyển sinh"""
        try:
            # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI TÌM KIẾM
            self.refresh_data_from_admin()

            # Tìm trong các category giáo dục
            education_categories = ['giáo dục vnexpress', 'giáo dục', 'thời sự vnexpress', 'thời sự thanh niên']
            all_videos = []

            for category in education_categories:
                videos = self.get_videos(category, 10)
                for video in videos:
                    video['search_category'] = category
                    all_videos.append(video)

            # Lọc video liên quan đến giáo dục
            education_videos = []
            for video in all_videos:
                title_lower = video['title'].lower()
                desc_lower = video.get('description', '').lower()

                # Kiểm tra từ khóa giáo dục
                edu_keywords = ['tuyển sinh', 'đại học', 'cao đẳng', 'giáo dục', 'học sinh', 'sinh viên', 'thi cử']
                if any(keyword in title_lower or keyword in desc_lower for keyword in edu_keywords):
                    education_videos.append(video)

            if education_videos:
                # Xếp hạng video giáo dục
                ranked_videos = self.rank_videos_by_relevance(education_videos, user_query)
                if ranked_videos:
                    best_video = ranked_videos[0]
                    self.current_video = best_video
                    self.current_category = best_video['search_category']
                    return self.create_article_from_video(best_video, user_query)

            return "❌ Không tìm thấy thông tin tuyển sinh phù hợp. Vui lòng thử lại với từ khóa cụ thể hơn."

        except Exception as e:
            return f"❌ Lỗi xử lý yêu cầu giáo dục: {str(e)}"

    def create_article_from_video(self, video, user_query):
        """Tạo bài báo từ video cụ thể với đánh giá phù hợp"""
        try:
            self.current_video = video
            self.current_category = video.get('search_category', 'thời sự')

            # Lấy transcript nếu là YouTube
            base_content = video['description']
            if video.get('source') == 'youtube' and video.get('video_id'):
                transcript = self.get_youtube_transcript(video['video_id'])
                if transcript:
                    base_content = transcript

            # Tạo nội dung bài báo
            article_content = self.content_generator.generate_article_content(
                video['title'],
                base_content,
                self.current_category,
                600
            )

            # Xử lý ảnh
            image_html = self.generate_image_html()

            # Đánh giá mức độ phù hợp SÂU
            relevance_evaluation = self.evaluate_deep_relevance(
                user_query,
                video['title'],
                video['description'],
                base_content,
                self.current_category  # Thêm video_category
            )

            # Tạo response
            icon = CATEGORY_ICONS.get(self.current_category, '📰')
            cleaned_title = self.content_generator.clean_title_for_content(video['title'])

            response = f"""## 🎯 KẾT QUẢ TÌM KIẾM

**🔍 Yêu cầu:** "{user_query}"
{relevance_evaluation}

## {icon} {cleaned_title}

**📅 Ngày đăng:** {video['published']}
**🔍 Nguồn:** {video.get('source', 'unknown').upper()}

{image_html}

{article_content}

---
**📺 Xem video gốc:** [{video['link']}]({video['link']})"""

            self.current_article = response
            return response

        except Exception as e:
            return f"❌ Lỗi tạo bài báo từ video: {str(e)}"

    def create_video_digest(self, category, user_query=None):
        """Tạo tổng hợp video với cấu trúc bài báo và ĐÁNH GIÁ ĐỘ PHÙ HỢP - ĐÃ SỬA"""
        try:
            # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI TẠO TỔNG HỢP
            self.refresh_data_from_admin()

            videos = self.get_videos(category, 5)
            if not videos:
                return f"❌ Không tìm thấy video nào trong chuyên mục {category}"

            # Lưu danh sách video tổng hợp
            self.digest_videos = videos

            icon = CATEGORY_ICONS.get(category, '📰')

            # Tạo cấu trúc bài báo cho tổng hợp video
            response = f"""## {icon} TỔNG HỢP {len(videos)} VIDEO {category.upper()}

### 📊 Tổng quan chuyên mục
Chuyên mục {category} trong 24 giờ qua đã có {len(videos)} video nổi bật với nhiều thông tin đáng chú ý. Dưới đây là tổng hợp chi tiết về các video quan trọng nhất.

"""

            for i, video in enumerate(videos, 1):
                source_badge = "🎬 YouTube" if video.get('source') == 'youtube' else "📡 RSS"

                # Xác định nguồn cụ thể
                source_detail = ""
                if video.get('source') == 'youtube':
                    if 'vtv24' in video.get('title', '').lower() or 'vtv24' in video.get('description', '').lower():
                        source_detail = " (VTV24)"
                    elif 'techrum' in video.get('title', '').lower() or 'techrum' in video.get('description', '').lower():
                        source_detail = " (Techrum)"
                else:
                    if 'vnexpress.net' in video.get('link', ''):
                        source_detail = " (VnExpress)"
                    elif 'thanhnien.vn' in video.get('link', ''):
                        source_detail = " (Thanh Niên)"
                    elif 'nld.com.vn' in video.get('link', ''):
                        source_detail = " (Người Lao Động)"
                    elif 'nguoiduatin.vn' in video.get('link', ''):
                        source_detail = " (Người Đưa Tin)"
                    elif 'tinmoi.vn' in video.get('link', ''):
                        source_detail = " (Tin Mới)"

                source_badge += source_detail

                # Làm sạch tiêu đề
                cleaned_title = self.content_generator.clean_title_for_content(video['title'])

                # Tạo nội dung bài báo cho từng video
                video_article = self.content_generator.generate_article_content(
                    cleaned_title,
                    video['description'],
                    category,
                    400  # Số từ cho mỗi video trong tổng hợp
                )

                # THÊM ĐÁNH GIÁ ĐỘ PHÙ HỢP - QUAN TRỌNG
                relevance_evaluation = self.evaluate_deep_relevance(
                    category,  # Dùng category làm query để đánh giá
                    video['title'],
                    video['description'],
                    video['description'],  # Dùng description làm nội dung
                    video.get('search_category', category)  # Thêm video_category
                )

                # Thumbnail nhỏ cho tổng hợp
                thumb_html = ""
                if video.get('thumbnail'):
                    thumb_html = f"<img src='{video['thumbnail']}' width='400' style='border-radius: 8px; margin: 10px 0;'>"
                else:
                    # Fallback thumbnail
                    thumb_html = f"<div style='width: 400px; height: 250px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 8px; display: flex; align-items: center; justify-content: center; color: white; font-size: 16px; font-weight: bold;'>🎬 {cleaned_title[:40]}...</div>"

                response += f"""
### {i}. {cleaned_title}

**Nguồn:** {source_badge}
**Thời gian đăng:** {video['published']}
**Độ dài ước tính:** 3-5 phút
**ĐỘ PHÙ HỢP:** {relevance_evaluation}

{thumb_html}

#### 📝 Tóm tắt nội dung
{video_article}

#### 🔗 Liên kết xem video
[🎬 Xem video gốc]({video['link']})

---

"""

            # Thêm phần kết luận cho tổng hợp
            response += f"""
### 🎯 Nhận định tổng quan
Các video trong chuyên mục {category} đã phản ánh đầy đủ những diễn biến mới nhất trong lĩnh vực này. Với {len(videos)} video chất lượng, người xem có cái nhìn toàn diện về các vấn đề đang được quan tâm.

**ĐÁNH GIÁ TỔNG THỂ:**
- 📊 **Số lượng video:** {len(videos)} video chất lượng
- 🎯 **Độ phù hợp chung:** Cao với chủ đề {category}
- 💡 **Khuyến nghị:** Người dùng nên xem video số 1 để nắm bắt thông tin quan trọng nhất, và video số 2-3 để hiểu sâu hơn về các khía cạnh liên quan.
"""

            return response

        except Exception as e:
            return f"❌ Lỗi tạo tổng hợp: {str(e)}"

    def select_video_from_digest(self, video_index):
        """Chọn video từ danh sách tổng hợp để tạo lời dẫn - ĐÃ SỬA LỖI QUAN TRỌNG"""
        try:
            if not self.digest_videos:
                return "❌ Không có danh sách video. Hãy tạo tổng hợp video trước."

            if video_index < 0 or video_index >= len(self.digest_videos):
                return f"❌ Số video phải từ 1 đến {len(self.digest_videos)}."

            selected_video = self.digest_videos[video_index]
            self.current_video = selected_video
            # THIẾT LẬP CATEGORY TỪ VIDEO ĐƯỢC CHỌN - QUAN TRỌNG
            self.current_category = selected_video.get('search_category', 'thời sự')
            
            # THÊM DÒNG DEBUG ĐỂ XÁC NHẬN
            print(f"ĐÃ CHỌN VIDEO {video_index+1}: {selected_video['title']}")
            
            return f"✅ Đã chọn Video {video_index+1}: {selected_video['title']}"

        except Exception as e:
            return f"❌ Lỗi khi chọn video: {str(e)}"

    def create_article_with_image(self, category, user_query=None):
        """Tạo bài báo với ảnh đầy đủ - ĐÃ SỬA LỖI VTV24 VÀ THÊM ĐÁNH GIÁ"""
        try:
            # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI TẠO BÀI BÁO
            self.refresh_data_from_admin()

            self.current_category = category

            # Lấy video - SỬA: TÌM TẤT CẢ CHỦ ĐỀ LIÊN QUAN
            videos = self.get_videos(category, 5)
            if not videos:
                return "❌ Không tìm thấy video phù hợp"

            # Chọn video đầu tiên
            self.current_video = videos[0]

            # Lấy transcript nếu là YouTube
            base_content = self.current_video['description']
            if self.current_video.get('source') == 'youtube' and self.current_video.get('video_id'):
                transcript = self.get_youtube_transcript(self.current_video['video_id'])
                if transcript:
                    base_content = transcript

            # Tạo nội dung bài báo theo cấu trúc mới (600+ từ)
            article_content = self.content_generator.generate_article_content(
                self.current_video['title'],
                base_content,
                category,
                600
            )

            # Xử lý ảnh - ĐẢM BẢO LUÔN CÓ ẢNH
            image_html = self.generate_image_html()

            # Đánh giá mức độ phù hợp - LUÔN HIỂN THỊ CHO MỌI BÀI BÁO
            relevance_evaluation = ""
            if user_query:
                relevance_evaluation = self.evaluate_deep_relevance(
                    user_query,
                    self.current_video['title'],
                    self.current_video['description'],
                    base_content,
                    self.current_category  # Thêm video_category
                )
            else:
                # Nếu không có user_query, vẫn hiển thị đánh giá dựa trên category
                relevance_evaluation = self.evaluate_deep_relevance(
                    category,
                    self.current_video['title'],
                    self.current_video['description'],
                    base_content,
                    self.current_category  # Thêm video_category
                )

            # Xác định nguồn cụ thể
            source_detail = ""
            if self.current_video.get('source') == 'youtube':
                source_badge = "🎬 YouTube"
                # Xác định kênh YouTube cụ thể nếu có thể
                if 'vtv24' in self.current_video.get('title', '').lower() or 'vtv24' in self.current_video.get('description', '').lower():
                    source_detail = " (VTV24)"
                elif 'techrum' in self.current_video.get('title', '').lower() or 'techrum' in self.current_video.get('description', '').lower():
                    source_detail = " (Techrum)"
                elif 'pdt' in self.current_video.get('title', '').lower() or 'pdt' in self.current_video.get('description', '').lower():
                    source_detail = " (PDT)"
                elif 'tin tức việt' in self.current_video.get('title', '').lower():
                    source_detail = " (Tin Tức Việt)"
                elif 'tin24h' in self.current_video.get('title', '').lower():
                    source_detail = " (Tin24h)"
                elif 'tin360' in self.current_video.get('title', '').lower():
                    source_detail = " (Tin360)"
            else:
                source_badge = "📡 RSS"
                # Xác định nguồn RSS cụ thể
                if 'vnexpress.net' in self.current_video.get('link', ''):
                    source_detail = " (VnExpress)"
                elif 'thanhnien.vn' in self.current_video.get('link', ''):
                    source_detail = " (Thanh Niên)"
                elif 'nld.com.vn' in self.current_video.get('link', ''):
                    source_detail = " (Người Lao Động)"
                elif 'nguoiduatin.vn' in self.current_video.get('link', ''):
                    source_detail = " (Người Đưa Tin)"
                elif 'tinmoi.vn' in self.current_video.get('link', ''):
                    source_detail = " (Tin Mới)"

            source_badge += source_detail

            # Tạo response hoàn chỉnh
            icon = CATEGORY_ICONS.get(category, '📰')

            # Làm sạch tiêu đề cho hiển thị
            cleaned_title = self.content_generator.clean_title_for_content(self.current_video['title'])

            response = f"""## {icon} {cleaned_title}

**📅 Ngày đăng:** {self.current_video['published']}
**🔍 Nguồn:** {source_badge}

{relevance_evaluation}

{image_html}

{article_content}

---
**📺 Xem video gốc:** [{self.current_video['link']}]({self.current_video['link']})"""

            self.current_article = response
            return response

        except Exception as e:
            print(f"❌ Lỗi tạo bài báo: {e}")
            return f"❌ Lỗi hệ thống: {str(e)}"

    def create_article_random_video(self, category, user_query=None):
        """Tạo bài báo từ video ngẫu nhiên trong cùng chuyên mục"""
        try:
            # LÀM MỚI DỮ LIỆU TỪ ADMIN TRƯỚC KHI TẠO BÀI BÁO
            self.refresh_data_from_admin()

            self.current_category = category

            # Lấy danh sách video - SỬA: TÌM TẤT CẢ CHỦ ĐỀ LIÊN QUAN
            videos = self.get_videos(category, 10)  # Lấy nhiều video để chọn ngẫu nhiên
            if not videos:
                return "❌ Không tìm thấy video phù hợp"

            # Chọn video ngẫu nhiên (không phải video đầu tiên)
            if len(videos) > 1:
                self.current_video = random.choice(videos[1:])
            else:
                self.current_video = videos[0]

            # Lấy transcript nếu là YouTube
            base_content = self.current_video['description']
            if self.current_video.get('source') == 'youtube' and self.current_video.get('video_id'):
                transcript = self.get_youtube_transcript(self.current_video['video_id'])
                if transcript:
                    base_content = transcript

            # Tạo nội dung bài báo
            article_content = self.content_generator.generate_article_content(
                self.current_video['title'],
                base_content,
                category,
                600
            )

            # Xử lý ảnh
            image_html = self.generate_image_html()

            # Đánh giá mức độ phù hợp - LUÔN HIỂN THỊ CHO MỌI BÀI BÁO
            relevance_evaluation = ""
            if user_query:
                relevance_evaluation = self.evaluate_deep_relevance(
                    user_query,
                    self.current_video['title'],
                    self.current_video['description'],
                    base_content,
                    self.current_category  # Thêm video_category
                )
            else:
                # Nếu không có user_query, vẫn hiển thị đánh giá dựa trên category
                relevance_evaluation = self.evaluate_deep_relevance(
                    category,
                    self.current_video['title'],
                    self.current_video['description'],
                    base_content,
                    self.current_category  # Thêm video_category
                )

            # Xác định nguồn cụ thể
            source_detail = ""
            if self.current_video.get('source') == 'youtube':
                source_badge = "🎬 YouTube"
                # Xác định kênh YouTube cụ thể nếu có thể
                if 'vtv24' in self.current_video.get('title', '').lower() or 'vtv24' in self.current_video.get('description', '').lower():
                    source_detail = " (VTV24)"
                elif 'techrum' in self.current_video.get('title', '').lower() or 'techrum' in self.current_video.get('description', '').lower():
                    source_detail = " (Techrum)"
                elif 'pdt' in self.current_video.get('title', '').lower() or 'pdt' in self.current_video.get('description', '').lower():
                    source_detail = " (PDT)"
                elif 'tin tức việt' in self.current_video.get('title', '').lower():
                    source_detail = " (Tin Tức Việt)"
                elif 'tin24h' in self.current_video.get('title', '').lower():
                    source_detail = " (Tin24h)"
                elif 'tin360' in self.current_video.get('title', '').lower():
                    source_detail = " (Tin360)"
            else:
                source_badge = "📡 RSS"
                # Xác định nguồn RSS cụ thể
                if 'vnexpress.net' in self.current_video.get('link', ''):
                    source_detail = " (VnExpress)"
                elif 'thanhnien.vn' in self.current_video.get('link', ''):
                    source_detail = " (Thanh Niên)"
                elif 'nld.com.vn' in self.current_video.get('link', ''):
                    source_detail = " (Người Lao Động)"
                elif 'nguoiduatin.vn' in self.current_video.get('link', ''):
                    source_detail = " (Người Đưa Tin)"
                elif 'tinmoi.vn' in self.current_video.get('link', ''):
                    source_detail = " (Tin Mới)"

            source_badge += source_detail

            # Tạo response
            icon = CATEGORY_ICONS.get(category, '📰')

            # Làm sạch tiêu đề cho hiển thị
            cleaned_title = self.content_generator.clean_title_for_content(self.current_video['title'])

            response = f"""## {icon} {cleaned_title} 🎲 (Video khác)

**📅 Ngày đăng:** {self.current_video['published']}
**🔍 Nguồn:** {source_badge}

{relevance_evaluation}

{image_html}

{article_content}

---
**📺 Xem video gốc:** [{self.current_video['link']}]({self.current_video['link']})"""

            self.current_article = response
            return response

        except Exception as e:
            print(f"❌ Lỗi tạo bài báo từ video khác: {e}")
            return f"❌ Lỗi hệ thống: {str(e)}"

    def generate_image_html(self):
        """Tạo HTML hiển thị ảnh - ĐÃ ĐẢM BẢO LUÔN CÓ ẢNH"""
        if not self.current_video:
            return ""

        image_source = "Hình ảnh từ video"

        # ƯU TIÊN 1: Thumbnail từ RSS/YouTube
        if self.current_video.get('thumbnail'):
            return f"""
<div style="text-align: center; margin: 20px 0;">
    <img src="{self.current_video['thumbnail']}" width="600" style="border-radius: 10px; box-shadow: 0 4px 8px rgba(0,0,0,0.1);">
    <p style="font-style: italic; color: #666; margin-top: 8px;">{image_source}: {self.content_generator.clean_title_for_content(self.current_video['title'])}</p>
</div>
"""

        # ƯU TIÊN 2: Ảnh trích xuất từ video
        elif self.current_video.get('extracted_frame') and os.path.exists(self.current_video['extracted_frame']):
            try:
                with open(self.current_video['extracted_frame'], "rb") as img_file:
                    img_base64 = base64.b64encode(img_file.read()).decode()
                return f"""
<div style="text-align: center; margin: 20px 0;">
    <img src="data:image/jpeg;base64,{img_base64}" width="600" style="border-radius: 10px; box-shadow: 0 4px 8px rgba(0,0,0,0.1);">
    <p style="font-style: italic; color: #666; margin-top: 8px;">{image_source} (trích xuất): {self.content_generator.clean_title_for_content(self.current_video['title'])}</p>
</div>
"""
            except Exception as e:
                print(f"❌ Lỗi hiển thị ảnh trích xuất: {e}")

        # ƯU TIÊN 3: Tạo ảnh fallback
        fallback_image = self.media_processor.create_fallback_image(
            self.current_video['title'],
            self.current_category or 'thời sự'
        )

        if fallback_image and os.path.exists(fallback_image):
            try:
                with open(fallback_image, "rb") as img_file:
                    img_base64 = base64.b64encode(img_file.read()).decode()
                return f"""
<div style="text-align: center; margin: 20px 0;">
    <img src="data:image/jpeg;base64,{img_base64}" width="600" style="border-radius: 10px; box-shadow: 0 4px 8px rgba(0,0,0,0.1);">
    <p style="font-style: italic; color: #666; margin-top: 8px;">{image_source} (minh họa): {self.content_generator.clean_title_for_content(self.current_video['title'])}</p>
</div>
"""
            except Exception as e:
                print(f"❌ Lỗi hiển thị ảnh fallback: {e}")

        # FALLBACK CUỐI CÙNG: Ảnh mặc định
        return f"""
<div style="text-align: center; margin: 20px 0;">
    <div style="width: 600px; height: 400px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 10px; display: flex; align-items: center; justify-content: center; color: white; font-size: 24px; font-weight: bold; box-shadow: 0 4px 8px rgba(0,0,0,0.1);">
        {CATEGORY_ICONS.get(self.current_category, '📰')} {self.content_generator.clean_title_for_content(self.current_video['title'])[:50]}...
    </div>
    <p style="font-style: italic; color: #666; margin-top: 8px;">{image_source}: {self.content_generator.clean_title_for_content(self.current_video['title'])}</p>
</div>
"""

    # ==============================================
    # CÁC PHƯƠNG THỨC TẠO SCRIPT - ĐÃ SỬA LỖI THIẾU PHƯƠNG THỨC
    # ==============================================

    def create_script(self, script_type="1 Cột"):
        """Tạo lời dẫn BTV với nội dung > 700 từ - ĐÃ SỬA LỖI LẶP ĐOẠN"""
        if not self.current_video:
            return "❌ Hãy chọn video trước khi tạo lời dẫn"

        try:
            title = self.content_generator.clean_title_for_content(self.current_video['title'])
            category = self.current_category or 'thời sự'

            # Tạo nội dung script chi tiết > 700 từ - ĐÃ SỬA LỖI LẶP
            script_content = self.content_generator.generate_detailed_script_content(title, category)

            if script_type == "1 Cột":
                script = self.create_one_column_script(title, category, script_content)
            elif script_type == "2 Cột":
                script = self.create_two_column_script(title, category, script_content)
            else:  # 3 Cột
                script = self.create_three_column_script(title, category, script_content)

            self.current_script = script
            return script

        except Exception as e:
            return f"❌ Lỗi tạo script: {str(e)}"

    def create_one_column_script(self, title, category, content):
        """Tạo script 1 cột với nội dung > 700 từ - ĐÃ SỬA LỖI LẶP"""
        return f"""# 🎤 LỜI DẪN BTV - {category.upper()}

**TIÊU ĐỀ:** {title}
**THỜI LƯỢNG:** 12-15 phút
**NGÀY PHÁT SÓNG:** {self.format_time()}
**BIÊN TẬP VIÊN:** [Tên BTV]

---

{content}

---

**KẾT THÚC CHƯƠNG TRÌNH**

Xin cảm ơn quý vị và các bạn đã theo dõi chương trình! Hẹn gặp lại trong những bản tin tiếp theo.
"""

    def create_two_column_script(self, title, category, content):
        """Tạo script 2 cột với cấu trúc phân đoạn"""
        # Phân chia nội dung thành các đoạn
        paragraphs = [p for p in content.split('\n\n') if p.strip()]

        table_rows = ""
        for i, para in enumerate(paragraphs, 1):
            if para.strip():
                # Làm sạch đoạn văn cho bảng
                clean_para = re.sub(r'#+\s*', '', para)  # Loại bỏ markdown headers
                table_rows += f"""
<tr>
    <td style="border: 1px solid #ddd; padding: 12px; width: 15%; font-weight: bold; vertical-align: top; background-color: #f8f9fa;">Đoạn {i}</td>
    <td style="border: 1px solid #ddd; padding: 12px; width: 85%; vertical-align: top;">{clean_para}</td>
</tr>
"""

        return f"""# 🎤 LỜI DẪN BTV - {category.upper()}

**TIÊU ĐỀ:** {title}
**ĐỊNH DẠNG:** 2 CỘT - PHÂN ĐOẠN CHI TIẾT
**THỜI LƯỢNG:** 12-15 phút
**BIÊN TẬP VIÊN:** [Tên BTV]

<table style="width: 100%; border-collapse: collapse; font-size: 14px; line-height: 1.6;">
{table_rows}
</table>

---
**KẾT THÚC CHƯƠNG TRÌNH**

Xin chân thành cảm ơn quý vị khán giả đã đồng hành cùng chúng tôi. Chúc quý vị một ngày làm việc hiệu quả!
"""

    def create_three_column_script(self, title, category, script_content):
        """Tạo script 3 cột với timeline và hướng dẫn"""
        # Tách nội dung thành các đoạn
        paragraphs = [p.strip() for p in script_content.split('\n\n') if p.strip()]

        # Tạo timeline 12-15 phút
        total_duration = 15 * 60  # 15 phút tính bằng giây
        segment_duration = total_duration // len(paragraphs) if paragraphs else 0

        table_rows = ""
        for i, para in enumerate(paragraphs):
            if para.strip():
                # Tính thời gian cho từng segment
                start_time = i * segment_duration
                end_time = (i + 1) * segment_duration

                start_min = start_time // 60
                start_sec = start_time % 60
                end_min = end_time // 60
                end_sec = end_time % 60

                time_range = f"{start_min:02d}:{start_sec:02d} - {end_min:02d}:{end_sec:02d}"

                # Làm sạch đoạn văn
                clean_para = re.sub(r'#+\s*', '', para)

                # Hướng dẫn phát thanh
                guidance = self.get_script_guidance(i, len(paragraphs))

                table_rows += f"""
<tr>
    <td style="border: 1px solid #ddd; padding: 10px; width: 15%; vertical-align: top; font-weight: bold; background-color: #f0f8ff;">{time_range}</td>
    <td style="border: 1px solid #ddd; padding: 10px; width: 60%; vertical-align: top;">{clean_para}</td>
    <td style="border: 1px solid #ddd; padding: 10px; width: 25%; vertical-align: top; font-size: 12px; color: #555; background-color: #fffaf0;">{guidance}</td>
</tr>
"""

        return f"""# 🎤 LỜI DẪN BTV - {category.upper()}

**TIÊU ĐỀ:** {title}
**ĐỊNH DẠNG:** 3 CỘT - TIMELINE CHI TIẾT
**THỜI LƯỢNG:** 12-15 phút
**BIÊN TẬP VIÊN:** [Tên BTV]

<table style="width: 100%; border-collapse: collapse; font-size: 13px; line-height: 1.5;">
{table_rows}
</table>

---
**KẾT THÚC CHƯƠNG TRÌNH**

Trân trọng cảm ơn quý vị đã theo dõi chương trình. Mọi ý kiến đóng góp xin gửi về hòm thư của đài!
"""

    def get_script_guidance(self, segment_index, total_segments):
        """Hướng dẫn phát thanh chi tiết"""
        guidance_map = {
            0: "🎙️ GIỌNG MỞ ĐẦU: Tươi vui, thân thiện, tạo thiện cảm. Tốc độ vừa phải, nhấn mạnh từ khóa quan trọng.",
            1: "🎙️ TRÌNH BÀY THÔNG TIN: Giọng rõ ràng, truyền cảm. Sử dụng ngữ điệu lên xuống để tạo điểm nhấn.",
            2: "🎙️ PHÂN TÍCH CHUYÊN SÂU: Giọng trầm ấm, thể hiện sự am hiểu. Tốc độ chậm rãi, nhấn mạnh thông tin quan trọng.",
            3: "🎙️ TRÌNH BÀY SỐ LIỆU: Giọng tự tin, rõ ràng từng con số. Tạm dừng sau các số liệu quan trọng.",
            4: "🎙️ PHẢN ỨNG VÀ ĐÁNH GIÁ: Giọng linh hoạt, thể hiện cảm xúc phù hợp. Sử dụng ngữ điệu biểu cảm.",
            5: "🎙️ GIẢI PHÁP VÀ KHUYẾN NGHỊ: Giọng tích cực, truyền cảm hứng. Nhấn mạnh các điểm then chốt.",
            6: "🎙️ CHUYỂN TIẾP CUỐI: Giọng ấm áp, tạo cảm xúc kết nối. Tốc độ chậm dần, tạo không khí kết thúc.",
            7: "🎙️ KẾT THÚC: Giọng chân thành, lưu luyến. Kết thúc với âm điệu tích cực, hẹn gặp lại."
        }

        return guidance_map.get(segment_index, "🎙️ TRÌNH BÀY THÔNG TIN: Giọng rõ ràng, tự tin. Duy trì tốc độ ổn định.")

    def export_content(self, format_type):
        """Xuất nội dung ra file - ĐÃ SỬA LỖI VÀ CẢI THIỆN"""
        try:
            if not self.current_article and not self.current_script:
                return None, "❌ Chưa có nội dung để xuất"

            # Sử dụng nội dung hiện tại (bài báo hoặc script)
            content = self.current_script if self.current_script else self.current_article
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # Đảm bảo thư mục tồn tại
            export_dir = "exported_files"
            os.makedirs(export_dir, exist_ok=True)

            if format_type == "TEXT":
                filename = f"noi_dung_{timestamp}.txt"
                filepath = os.path.join(export_dir, filename)

                # Làm sạch nội dung cho text
                clean_content = re.sub(r'<[^>]+>', '', content)
                clean_content = re.sub(r'!\[.*?\]\(.*?\)', '', content)
                clean_content = re.sub(r'#{1,6}\s*', '', clean_content)  # Loại bỏ markdown headers

                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(clean_content)

                return filepath, f"✅ Đã xuất file TEXT: {filename}"

            elif format_type == "DOC":
                filename = f"noi_dung_{timestamp}.docx"
                filepath = os.path.join(export_dir, filename)

                doc = Document()
                doc.add_heading('Nội dung xuất từ Chatbot', 0)

                # Thêm thông tin metadata
                if self.current_video:
                    doc.add_paragraph(f"Tiêu đề: {self.current_video.get('title', 'Không có tiêu đề')}")
                    doc.add_paragraph(f"Nguồn: {self.current_video.get('source', 'Không xác định')}")
                    doc.add_paragraph(f"Thời gian: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

                # Thêm nội dung
                clean_content = re.sub(r'<[^>]+>', '', content)
                paragraphs = clean_content.split('\n')

                for para in paragraphs:
                    if para.strip():
                        if para.startswith('##'):
                            doc.add_heading(para.replace('##', '').strip(), 2)
                        elif para.startswith('###'):
                            doc.add_heading(para.replace('###', '').strip(), 3)
                        else:
                            # Kiểm tra nếu là danh sách
                            if para.strip().startswith('-') or para.strip().startswith('•'):
                                p = doc.add_paragraph()
                                p.add_run(para.strip()).bold = False
                            else:
                                doc.add_paragraph(para.strip())

                doc.save(filepath)
                return filepath, f"✅ Đã xuất file DOCX: {filename}"

            else:  # PDF
                filename = f"noi_dung_{timestamp}.pdf"
                filepath = os.path.join(export_dir, filename)

                # Tạo HTML tạm với CSS cải thiện
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <title>Nội dung xuất từ Chatbot</title>
                    <style>
                        body {{
                            font-family: 'DejaVu Sans', Arial, sans-serif;
                            line-height: 1.6;
                            margin: 40px;
                            color: #333;
                        }}
                        h1 {{
                            color: #2c3e50;
                            border-bottom: 3px solid #3498db;
                            padding-bottom: 10px;
                            text-align: center;
                        }}
                        h2 {{
                            color: #34495e;
                            border-left: 4px solid #3498db;
                            padding-left: 10px;
                            margin-top: 30px;
                        }}
                        h3 {{
                            color: #7f8c8d;
                            margin-top: 20px;
                        }}
                        .content {{
                            margin: 20px 0;
                            text-align: justify;
                        }}
                        .metadata {{
                            background-color: #f8f9fa;
                            padding: 15px;
                            border-radius: 5px;
                            margin: 20px 0;
                            border-left: 4px solid #2ecc71;
                        }}
                        table {{
                            width: 100%;
                            border-collapse: collapse;
                            margin: 10px 0;
                            font-size: 14px;
                        }}
                        th, td {{
                            border: 1px solid #ddd;
                            padding: 12px;
                            text-align: left;
                        }}
                        th {{
                            background-color: #f2f2f2;
                            font-weight: bold;
                        }}
                        .footer {{
                            margin-top: 40px;
                            padding-top: 20px;
                            border-top: 1px solid #ddd;
                            text-align: center;
                            color: #7f8c8d;
                            font-size: 12px;
                        }}
                    </style>
                </head>
                <body>
                <h1>Nội dung xuất từ Chatbot</h1>

                <div class="metadata">
                    <strong>Thời gian xuất:</strong> {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}<br>
                    {f"<strong>Tiêu đề video:</strong> {self.current_video.get('title', 'Không có tiêu đề')}<br>" if self.current_video else ""}
                    {f"<strong>Nguồn:</strong> {self.current_video.get('source', 'Không xác định')}<br>" if self.current_video else ""}
                    {f"<strong>Chuyên mục:</strong> {self.current_category}<br>" if self.current_category else ""}
                </div>

                <div class="content">{content}</div>

                <div class="footer">
                    Được tạo bởi Chatbot Tin Tức Video Thông Minh<br>
                    {datetime.now().strftime('Ngày %d/%m/%Y lúc %H:%M:%S')}
                </div>
                </body>
                </html>
                """

                try:
                    # Cấu hình pdfkit với options chi tiết
                    options = {
                        'page-size': 'A4',
                        'margin-top': '1.0in',
                        'margin-right': '0.75in',
                        'margin-bottom': '1.0in',
                        'margin-left': '0.75in',
                        'encoding': "UTF-8",
                        'no-outline': None,
                        'enable-local-file-access': None,
                        'footer-center': f'Trang [page] / [topage] - {datetime.now().strftime("%d/%m/%Y")}',
                        'footer-font-size': '10',
                        'footer-font-name': 'DejaVu Sans'
                    }

                    # Thử tạo PDF
                    pdfkit.from_string(html_content, filepath, options=options)
                    return filepath, f"✅ Đã xuất file PDF: {filename}"

                except Exception as pdf_error:
                    print(f"❌ Lỗi tạo PDF: {pdf_error}")
                    # Fallback to text nếu lỗi PDF
                    return self.export_content("TEXT")

        except Exception as e:
            print(f"❌ Lỗi xuất file: {e}")
            return None, f"❌ Lỗi xuất file: {str(e)}"

    # ==============================================
    # PHƯƠNG THỨC QUẢN LÝ LỊCH SỬ - MỚI THÊM
    # ==============================================

    def add_to_history(self, user_message: str, assistant_response: str, metadata: Dict = None):
        """Thêm tin nhắn vào lịch sử"""
        if self.history_manager.current_session_id:
            self.history_manager.add_message(
                self.history_manager.current_session_id,
                'user',
                user_message,
                metadata
            )
            self.history_manager.add_message(
                self.history_manager.current_session_id,
                'assistant',
                assistant_response,
                metadata
            )

    def create_new_chat_session(self, title: str = "Cuộc trò chuyện mới"):
        """Tạo phiên chat mới"""
        return self.history_manager.create_new_session(title)

    def get_chat_sessions(self):
        """Lấy danh sách tất cả phiên chat"""
        return self.history_manager.get_all_sessions()

    def load_session(self, session_id: str):
        """Tải phiên chat cụ thể"""
        session = self.history_manager.get_session(session_id)
        if session:
            self.history_manager.current_session_id = session_id
            return session['messages']
        return []

    def delete_session(self, session_id: str):
        """Xóa phiên chat"""
        self.history_manager.delete_session(session_id)

    def clear_all_history(self):
        """Xóa toàn bộ lịch sử"""
        self.history_manager.clear_all_sessions()
        self.create_new_chat_session("Phiên làm việc mới")

print("✅ Backend đã được khởi tạo thành công!")