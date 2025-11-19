import os
import shutil
import uuid
import json
from docx import Document
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import FileResponse, RedirectResponse

router = APIRouter(prefix="/api/clusters", tags=["Clusters"])

DATA_DIR = "data"
os.makedirs(DATA_DIR, exist_ok=True)

CLUSTERS_PATH = os.path.join(DATA_DIR, "latest_clusters.json")

@app.get("/api/claims-summary")
def get_claims_summary():
    file_path = os.path.join("backend", "data", "claims_grouped_summary.json")
    return FileResponse(file_path, media_type="application/json")

def load_clusters():
    if os.path.exists(CLUSTERS_PATH):
        with open(CLUSTERS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_clusters(clusters):
    with open(CLUSTERS_PATH, "w", encoding="utf-8") as f:
        json.dump(clusters, f, ensure_ascii=False, indent=2)

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    fn = file.filename
    if not fn.endswith((".json", ".csv")):
        raise HTTPException(status_code=400, detail="Only .json or .csv allowed")
    dest = os.path.join(DATA_DIR, "crawler_output" + os.path.splitext(fn)[1])
    with open(dest, "wb") as out:
        shutil.copyfileobj(file.file, out)
    # Run the processing pipeline
    clusters = cluster_articles_only(dest, output_prefix=os.path.join(DATA_DIR, "news_ai"))
    # annotate clusters with status and id
    for c in clusters:
        c.setdefault("status", "Draft")
    save_clusters(clusters)
    return {"message": "File processed", "n_clusters": len(clusters)}

@app.get("/api/clusters")
def get_clusters():
    clusters = load_clusters()
    # return light-weight list
    return [ { "cluster_id": c["cluster_id"], "n_articles": c["n_articles"], "summary": c["summary"], "script": c["script"], "status": c.get("status","Draft"), "articles": c["articles"] } for c in clusters ]

@app.get("/api/clusters/{cluster_id}")
def get_cluster(cluster_id: int):
    clusters = load_clusters()
    for c in clusters:
        if int(c["cluster_id"]) == int(cluster_id):
            return c
    raise HTTPException(status_code=404, detail="Cluster not found")

@app.post("/api/clusters/{cluster_id}/script")
def save_script(cluster_id: int, payload: dict):
    clusters = load_clusters()
    for c in clusters:
        if int(c["cluster_id"]) == int(cluster_id):
            c["script"] = payload.get("script", c.get("script",""))
            save_clusters(clusters)
            return {"message": "Script saved"}
    raise HTTPException(status_code=404, detail="Cluster not found")

@app.post("/api/clusters/{cluster_id}/export")
def export_cluster(cluster_id: int):
    clusters = load_clusters()
    cluster = next((c for c in clusters if int(c["cluster_id"]) == int(cluster_id)), None)
    if not cluster:
        raise HTTPException(status_code=404, detail="Cluster not found")

    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=500, detail="Missing dependency: python-docx")

    doc = Document()
    doc.add_heading(f"Cluster {cluster_id}", level=1)
    doc.add_paragraph(f"Status: {cluster['status']}")
    doc.add_paragraph(f"Summary: {cluster['summary']}")
    doc.add_paragraph("")

    doc.add_heading("Script", level=2)
    doc.add_paragraph(cluster.get("script", ""))
    doc.add_paragraph("")

    doc.add_heading("Articles", level=2)
    for a in cluster.get("articles", []):
        doc.add_paragraph(f"Title: {a.get('Title', '')}")
        doc.add_paragraph(f"Source: {a.get('URL', '')}")
        doc.add_paragraph(f"Date: {a.get('Date', '')}")
        doc.add_paragraph(a.get("Summary", ""))
        doc.add_paragraph("---")

    out_path = os.path.join(DATA_DIR, f"cluster_{cluster_id}.docx")
    doc.save(out_path)
    return FileResponse(out_path, filename=f"cluster_{cluster_id}.docx")

@app.post("/api/clusters/{cluster_id}/generate")
def generate_cluster_script(cluster_id: int):
    try:
        from article_cluster_pipeline import summarize_and_generate
        cluster = summarize_and_generate(cluster_id, data_dir=DATA_DIR, prefix="news_ai")

        # ✅ Properly update the in-memory cluster list and save it
        clusters = load_clusters()
        for i, c in enumerate(clusters):
            if int(c["cluster_id"]) == int(cluster_id):
                clusters[i] = cluster
                break
        save_clusters(clusters)

        print(f"✅ Updated cluster {cluster_id} with new summary/script.")
        return {"message": "Script generated", "cluster": cluster}

    except Exception as e:
        print("❌ Error generating script:", e)
        raise HTTPException(status_code=500, detail=str(e))